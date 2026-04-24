"""Template fill service - writes extracted data back to templates of any format."""
from __future__ import annotations

import csv
import os
import re
from copy import copy
from typing import Any

from ..core.exceptions import FillError
from ..core.logging import logger
from ..schemas.models import CandidateEvidence, FilledFieldResult, FilledResult, TemplateSchema
from ..utils.entity_utils import describe_entity_reason, entity_header_kind, evaluate_entity_compatibility, validate_entity_value
from ..utils.text_utils import clean_cell_value


PLACEHOLDER_PATTERNS = [r'\{\{([^}]+)\}\}', r'\{([^}]+)\}', r'【([^】]+)】']


def fill_template(template: TemplateSchema, extracted_data: list[dict], output_path: str) -> FilledResult:
    """Fill a template with extracted data and save to output_path."""
    file_type = template.file_type
    logger.info("Filling template: %s (type=%s) -> %s", template.source_file, file_type, output_path)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    result = FilledResult(
        template_file=template.source_file,
        output_file=output_path,
        entity_legality_report={
            "validated_fields": 0,
            "blocked_count": 0,
            "dropped_records": 0,
            "blocked_examples": [],
            "accepted_examples": [],
            "filtered_records": [],
            "filter_reason_counts": {},
            "per_source_filter_reasons": {},
            "recovered_examples": [],
        },
    )

    try:
        if file_type == "excel":
            _fill_excel(template, extracted_data, output_path, result)
        elif file_type == "word":
            _fill_word(template, extracted_data, output_path, result)
        elif file_type == "markdown":
            _fill_markdown(template, extracted_data, output_path, result)
        elif file_type == "text":
            _fill_text(template, extracted_data, output_path, result)
        elif file_type == "csv":
            _fill_csv(template, extracted_data, output_path, result)
        else:
            raise FillError(f"Unsupported template type: {file_type}")
    except FillError:
        raise
    except Exception as exc:
        raise FillError(f"Error filling template: {exc}") from exc

    _merge_table_reports(result, extracted_data)
    result.metric_definitions = _metric_definitions()
    # Compute intermediate fill_rate so the log message is meaningful.
    # (pipeline_service._finalize_result_metrics will recompute with expected_rows later.)
    if result.rows_filled > 0:
        result.fill_rate = 100.0  # provisional; pipeline normalizes to rows_filled/expected_rows
    else:
        total_fields = len(result.filled_fields)
        filled_fields = sum(1 for field in result.filled_fields if field.value not in (None, "", "N/A"))
        result.fill_rate = (filled_fields / total_fields * 100) if total_fields > 0 else 0.0
    logger.info("  -> Fill rate: %.1f%%, %s rows", result.fill_rate, result.rows_filled)
    return result


def _fill_excel(template: TemplateSchema, extracted_data: list[dict], output_path: str, result: FilledResult):
    """Fill an Excel template."""
    from openpyxl import load_workbook

    try:
        workbook = load_workbook(template.source_file)
    except Exception:
        workbook = None

    if workbook is not None:
        _fill_excel_openpyxl(workbook, template, extracted_data, output_path, result)
    else:
        _fill_excel_pandas(template, extracted_data, output_path, result)


def _fill_excel_openpyxl(workbook, template: TemplateSchema, extracted_data: list[dict], output_path: str, result: FilledResult):
    """Fill Excel while preserving formatting when possible."""
    for table_data in extracted_data:
        table_index = table_data.get("table_index", 0)
        template_table = next((item for item in template.tables if item.table_index == table_index), None)
        if template_table is None:
            continue

        sheet_name = template_table.sheet_name
        if sheet_name and sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
        elif workbook.sheetnames:
            worksheet = workbook[workbook.sheetnames[min(table_index, len(workbook.sheetnames) - 1)]]
            sheet_name = worksheet.title
        else:
            continue

        headers = table_data.get("headers", [])
        records = _prepare_records_for_writeback(table_data, result)
        header_row = _find_excel_header_row(worksheet, headers)
        column_map = _map_headers_to_sheet(worksheet, header_row, headers)
        data_start_row = header_row + 1
        writable_rows = sorted(set(template_table.writable_rows)) if template_table and template_table.writable_rows else []
        target_rows = [data_start_row + row_offset for row_offset in writable_rows]
        append_base_row = max(worksheet.max_row, target_rows[-1] if target_rows else data_start_row - 1)
        appended_rows = 0

        if len(records) > len(target_rows) and target_rows:
            result.warnings.append(
                f"Excel 模板表 {table_index} 预留 {len(target_rows)} 个可写行位，已为额外 {len(records) - len(target_rows)} 条记录追加新行"
            )

        for row_offset, record in enumerate(records):
            if row_offset < len(target_rows):
                worksheet_row = target_rows[row_offset]
            else:
                worksheet_row = append_base_row + appended_rows + 1
                style_row = target_rows[-1] if target_rows else max(header_row + 1, worksheet_row - 1)
                _clone_excel_row_style(worksheet, style_row, worksheet_row)
                appended_rows += 1
            for col_index, header in enumerate(headers):
                if col_index not in column_map:
                    continue
                worksheet_col = column_map[col_index]
                cell = worksheet.cell(worksheet_row, worksheet_col)
                value = record["values"].get(header, "")
                normalized = _smart_value(value)
                cell.value = normalized
                result.filled_fields.append(_build_field_result(
                    header,
                    f"{sheet_name}!{cell.coordinate}",
                    value,
                    normalized,
                    record,
                    table_data,
                ))

        result.rows_filled += len(records)
        result.record_count += len(records)

    workbook.save(output_path)


def _clone_excel_row_style(worksheet, source_row: int, target_row: int):
    """Copy row height and cell style so appended rows keep the template structure."""
    if source_row <= 0 or source_row == target_row:
        return
    if source_row > worksheet.max_row:
        return
    worksheet.row_dimensions[target_row].height = worksheet.row_dimensions[source_row].height
    for column in range(1, worksheet.max_column + 1):
        source_cell = worksheet.cell(source_row, column)
        target_cell = worksheet.cell(target_row, column)
        if source_cell.has_style:
            target_cell._style = copy(source_cell._style)


def _fill_excel_pandas(template: TemplateSchema, extracted_data: list[dict], output_path: str, result: FilledResult):
    """Fallback Excel writer when openpyxl cannot load the template."""
    import pandas as pd

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        for table_data in extracted_data:
            headers = table_data.get("headers", [])
            records = _prepare_records_for_writeback(table_data, result)
            table_index = table_data.get("table_index", 0)
            template_table = next((item for item in template.tables if item.table_index == table_index), None)
            sheet_name = template_table.sheet_name if template_table and template_table.sheet_name else f"Sheet{table_index + 1}"
            frame = pd.DataFrame([
                [record["values"].get(header, "") for header in headers]
                for record in records
            ], columns=headers)
            frame.to_excel(writer, sheet_name=sheet_name[:31], index=False)

            for row_offset, record in enumerate(records, start=2):
                for col_index, header in enumerate(headers, start=1):
                    value = record["values"].get(header, "")
                    result.filled_fields.append(_build_field_result(
                        header,
                        f"{sheet_name}!R{row_offset}C{col_index}",
                        value,
                        value,
                        record,
                        table_data,
                    ))
            result.rows_filled += len(records)
            result.record_count += len(records)


def _fill_word(template: TemplateSchema, extracted_data: list[dict], output_path: str, result: FilledResult):
    """Fill a Word template."""
    from docx import Document

    document = Document(template.source_file)
    for table_data in extracted_data:
        table_index = table_data.get("table_index", 0)
        if table_index >= len(document.tables):
            continue

        headers = table_data.get("headers", [])
        records = _prepare_records_for_writeback(table_data, result)
        template_table = next((item for item in template.tables if item.table_index == table_index), None)
        word_table = document.tables[table_index]
        word_headers = [cell.text.strip() for cell in word_table.rows[0].cells]
        column_map = _map_headers_to_row(headers, word_headers)
        writable_rows = list(template_table.writable_rows) if template_table and template_table.writable_rows else []
        if writable_rows:
            writable_row_indices = [row_index + 1 for row_index in writable_rows]
        else:
            writable_row_indices = list(range(1, len(word_table.rows)))
        if not writable_row_indices and records:
            writable_row_indices = list(range(1, len(word_table.rows)))

        if len(records) > len(writable_row_indices) and writable_row_indices:
            extra_count = len(records) - len(writable_row_indices)
            result.warnings.append(
                f"Word 模板表 {table_index} 预留 {len(writable_row_indices)} 个可写行位，已为额外 {extra_count} 条记录自动扩展表格行"
            )

        for record_index, record in enumerate(records):
            if record_index < len(writable_row_indices):
                row_offset = writable_row_indices[record_index]
            else:
                # Auto-expand: add new rows to the Word table
                style_row_idx = writable_row_indices[-1] if writable_row_indices else max(0, len(word_table.rows) - 1)
                row_offset = len(word_table.rows)
                _add_word_table_row(word_table, style_row_idx)
            if row_offset >= len(word_table.rows):
                continue
            word_row = word_table.rows[row_offset]
            for col_index, header in enumerate(headers):
                if col_index not in column_map:
                    continue
                word_col = column_map[col_index]
                if word_col >= len(word_row.cells):
                    continue
                value = record["values"].get(header, "")
                _set_word_cell_text(word_row.cells[word_col], str(value) if value else "")
                result.filled_fields.append(_build_field_result(
                    header,
                    f"table{table_index}.row{row_offset}.col{word_col}",
                    value,
                    value,
                    record,
                    table_data,
                ))

        result.rows_filled += len(records)
        result.record_count += len(records)

    _replace_word_placeholders(document, extracted_data, result)
    document.save(output_path)


def _replace_word_placeholders(document, extracted_data: list[dict], result: FilledResult):
    """Replace placeholders like {{field}} in Word paragraphs."""
    placeholder_values = _collect_placeholder_values(extracted_data)
    for paragraph in _iter_word_paragraphs(document):
        text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        updated = _replace_placeholders_in_text(text, placeholder_values)
        if updated == text:
            continue
        _set_word_paragraph_text(paragraph, updated, placeholder_values)

    for field_name, value in placeholder_values.items():
        if not value:
            continue
        result.filled_fields.append(FilledFieldResult(
            field_name=field_name,
            target_location="word_placeholder",
            value=value,
            normalized_value=value,
            confidence=_placeholder_confidence(field_name, extracted_data),
            evidence=_placeholder_evidence(field_name, extracted_data),
            source_file=_placeholder_source_file(field_name, extracted_data),
            match_method="placeholder_replace",
        ))


def _iter_word_paragraphs(document):
    """Yield all document paragraphs, including paragraphs inside table cells."""
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _add_word_table_row(word_table, style_row_idx: int):
    """Add a new row to a Word table, copying cell widths and paragraph formatting from style_row_idx."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = word_table._tbl
    # Clone the style row's XML as a template
    if style_row_idx < len(word_table.rows):
        style_row = word_table.rows[style_row_idx]
        new_tr = copy(style_row._tr)
        # Clear text content in all cells of the cloned row
        for tc in new_tr.findall(qn('w:tc')):
            for p in tc.findall(qn('w:p')):
                for r in p.findall(qn('w:r')):
                    for t in r.findall(qn('w:t')):
                        t.text = ''
        tbl.append(new_tr)
    else:
        word_table.add_row()


def _set_word_cell_text(cell, text: str):
    """Update a Word cell while preserving paragraph and first-run styling when possible."""
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
    else:
        paragraph = cell.add_paragraph()
    _set_word_paragraph_text(paragraph, text, {})
    for extra_paragraph in cell.paragraphs[1:]:
        _set_word_paragraph_text(extra_paragraph, "", {})


def _set_word_paragraph_text(paragraph, text: str, placeholder_values: dict[str, str]):
    """Replace paragraph text with a formatting-preserving best effort."""
    if paragraph.runs:
        changed_in_runs = False
        if placeholder_values:
            for run in paragraph.runs:
                updated_run = _replace_placeholders_in_text(run.text, placeholder_values)
                if updated_run != run.text:
                    run.text = updated_run
                    changed_in_runs = True
        if changed_in_runs:
            current_text = "".join(run.text for run in paragraph.runs)
            if current_text == text:
                return
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.text = text


def _fill_markdown(template: TemplateSchema, extracted_data: list[dict], output_path: str, result: FilledResult):
    """Fill a Markdown template."""
    content = template.raw_text
    for table_data in extracted_data:
        headers = table_data.get("headers", [])
        records = _prepare_records_for_writeback(table_data, result)
        table_index = table_data.get("table_index", 0)
        if not headers:
            continue

        markdown_table = _build_markdown_table(headers, records)
        table_pattern = r'(\|.+\|\n\|[\s:|-]+\|\n(?:\|.+\|\n?)*)'
        tables = list(re.finditer(table_pattern, content))
        if table_index < len(tables):
            match = tables[table_index]
            content = content[:match.start()] + markdown_table + content[match.end():]
        else:
            content += "\n\n" + markdown_table

        for row_offset, record in enumerate(records, start=1):
            for col_index, header in enumerate(headers, start=1):
                value = record["values"].get(header, "")
                result.filled_fields.append(_build_field_result(
                    header,
                    f"md_table{table_index}.row{row_offset}.col{col_index}",
                    value,
                    value,
                    record,
                    table_data,
                ))
        result.rows_filled += len(records)
        result.record_count += len(records)

    content = _replace_placeholders_in_text(content, _collect_placeholder_values(extracted_data))
    with open(output_path, "w", encoding="utf-8") as handle:
        handle.write(content)


def _fill_text(template: TemplateSchema, extracted_data: list[dict], output_path: str, result: FilledResult):
    """Fill a text template, replacing only the table region and placeholders."""
    content = template.raw_text
    if template.tables:
        for table_index, table_data in enumerate(extracted_data):
            headers = table_data.get("headers", [])
            records = _prepare_records_for_writeback(table_data, result)
            if not headers:
                continue
            delimiter = "\t" if "\t" in content or "," not in content else ","
            new_table_lines = [delimiter.join(headers)]
            for row_offset, record in enumerate(records, start=1):
                values = [str(record["values"].get(header, "")) for header in headers]
                new_table_lines.append(delimiter.join(values))
                for col_index, header in enumerate(headers, start=1):
                    value = record["values"].get(header, "")
                    result.filled_fields.append(_build_field_result(
                        header,
                        f"text_table{table_index}.row{row_offset}.col{col_index}",
                        value,
                        value,
                        record,
                        table_data,
                    ))
            new_table_text = "\n".join(new_table_lines)
            # Try to locate the original table region and replace only that part
            content = _replace_text_table_region(content, headers, delimiter, new_table_text)
            result.rows_filled += len(records)
            result.record_count += len(records)

    content = _replace_placeholders_in_text(content, _collect_placeholder_values(extracted_data))
    with open(output_path, "w", encoding="utf-8") as handle:
        handle.write(content)


def _replace_text_table_region(content: str, headers: list[str], delimiter: str, new_table_text: str) -> str:
    """Find the table region in the original content and replace only that region."""
    lines = content.split('\n')
    header_line = delimiter.join(headers)
    # Find the header line in the original content
    table_start = -1
    table_end = -1
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        # Check if this line matches the header
        line_fields = [f.strip() for f in stripped.split(delimiter)]
        if len(line_fields) >= len(headers) * 0.6:
            matching = sum(1 for h in headers if any(h in f or f in h for f in line_fields))
            if matching >= len(headers) * 0.5:
                table_start = i
                # Find the end of the table (contiguous non-empty lines with delimiter)
                table_end = i + 1
                for j in range(i + 1, len(lines)):
                    stripped_j = lines[j].strip()
                    if not stripped_j:
                        # Allow one blank line within table
                        if j + 1 < len(lines) and lines[j + 1].strip() and delimiter in lines[j + 1]:
                            table_end = j + 1
                            continue
                        break
                    if delimiter in stripped_j or not stripped_j:
                        table_end = j + 1
                    else:
                        break
                break

    if table_start >= 0:
        # Replace only the table region, preserve everything before and after
        before = '\n'.join(lines[:table_start])
        after = '\n'.join(lines[table_end:])
        parts = [p for p in [before, new_table_text, after] if p]
        return '\n'.join(parts)
    else:
        # Fallback: couldn't find the table, append it
        return content + "\n\n" + new_table_text


def _fill_csv(template: TemplateSchema, extracted_data: list[dict], output_path: str, result: FilledResult):
    """Fill a CSV template."""
    for table_data in extracted_data:
        headers = table_data.get("headers", [])
        records = _prepare_records_for_writeback(table_data, result)
        if not headers:
            continue
        with open(output_path, "w", newline="", encoding="utf-8") as handle:
            writer = csv.writer(handle)
            writer.writerow(headers)
            for row_offset, record in enumerate(records, start=1):
                row = [record["values"].get(header, "") for header in headers]
                writer.writerow(row)
                for col_index, header in enumerate(headers, start=1):
                    value = record["values"].get(header, "")
                    result.filled_fields.append(_build_field_result(
                        header,
                        f"csv.row{row_offset}.col{col_index}",
                        value,
                        value,
                        record,
                        table_data,
                    ))
        result.rows_filled += len(records)
        result.record_count += len(records)
        break


def _table_records(table_data: dict) -> list[dict]:
    """Normalize table data to record-shaped objects."""
    records = table_data.get("records") or []
    if records:
        return records

    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    return [
        {
            "values": {header: clean_cell_value(row[index]) if index < len(row) else "" for index, header in enumerate(headers)},
            "field_confidence": {},
            "field_evidence": {},
            "source_file": "",
            "match_methods": {},
        }
        for row in rows
    ]


def _prepare_records_for_writeback(table_data: dict, result: FilledResult) -> list[dict]:
    """Apply entity-legality checks before any template cells are written."""
    records = _non_empty_records(_table_records(table_data))
    headers = table_data.get("headers", [])
    entity_headers = [header for header in headers if entity_header_kind(header)]
    if not records or not entity_headers:
        return records

    report = result.entity_legality_report or {
        "validated_fields": 0,
        "blocked_count": 0,
        "dropped_records": 0,
        "blocked_examples": [],
        "accepted_examples": [],
        "filtered_records": [],
        "filter_reason_counts": {},
        "per_source_filter_reasons": {},
        "recovered_examples": [],
    }
    context = table_data.get("template_context") or {}
    context_text = " ".join(
        str(item) for item in (
            context.get("template_file", ""),
            context.get("anchor_text", ""),
            context.get("topic_text", ""),
        )
        if item
    )
    cleaned_records: list[dict] = []

    for record in records:
        values = record.setdefault("values", {})
        field_confidence = record.setdefault("field_confidence", {})
        field_evidence = record.setdefault("field_evidence", {})
        match_methods = record.setdefault("match_methods", {})

        for header in entity_headers:
            raw_value = clean_cell_value(values.get(header, ""))
            if not raw_value:
                continue
            report["validated_fields"] = int(report.get("validated_fields", 0)) + 1
            assessment = record.get("entity_compatibility", {}).get(header) or evaluate_entity_compatibility(
                raw_value,
                header,
                peer_headers=headers,
                context_text=context_text,
                record_values=values,
            )
            # Sub-national entities preserved by the extraction pipeline (e.g., province
            # records from a source with multiple distinct provinces) carry
            # recoverable_mismatch_reason="sub_national_preserved" in their
            # entity_compatibility metadata.  Treat them as accepted at write-back time
            # so their entity field is not erased and the record is not dropped.
            if not assessment.get("accepted") and assessment.get("recoverable_mismatch_reason") == "sub_national_preserved":
                assessment = dict(assessment)
                assessment["accepted"] = True
            if assessment.get("accepted"):
                _append_entity_example(
                    report,
                    "accepted_examples",
                    {
                        "field_name": header,
                        "value": raw_value,
                        "source_file": record.get("source_file", ""),
                    },
                )
                if assessment.get("recoverable_mismatch_reason") or assessment.get("compatibility_score", 1.0) < 0.95:
                    _append_entity_example(
                        report,
                        "recovered_examples",
                        {
                            "field_name": header,
                            "value": raw_value,
                            "reason": describe_entity_reason(
                                str(assessment.get("recoverable_mismatch_reason") or assessment.get("filter_reason") or "")
                            ),
                            "source_file": record.get("source_file", ""),
                        },
                    )
                continue

            reason = str(assessment.get("filter_reason") or assessment.get("recoverable_mismatch_reason") or "entity_incompatible")
            values[header] = ""
            field_confidence[header] = None
            field_evidence[header] = []
            match_methods[header] = "entity_legality_blocked"
            report["blocked_count"] = int(report.get("blocked_count", 0)) + 1
            report.setdefault("filter_reason_counts", {})
            report["filter_reason_counts"][reason] = int(report["filter_reason_counts"].get(reason, 0)) + 1
            source_file = record.get("source_file", "")
            if source_file:
                per_source = report.setdefault("per_source_filter_reasons", {}).setdefault(source_file, {})
                per_source[reason] = int(per_source.get(reason, 0)) + 1
            report.setdefault("filtered_records", []).append({
                "record_id": record.get("record_id", ""),
                "source_file": source_file,
                "field_name": header,
                "entity_text": raw_value,
                "normalized_entity_type": assessment.get("normalized_entity_type", ""),
                "normalized_granularity": assessment.get("normalized_granularity", ""),
                "filter_reason": reason,
                "filter_stage": assessment.get("filter_stage", "entity_compatibility"),
                "whether_recoverable": bool(assessment.get("whether_recoverable", False)),
            })
            _append_entity_example(
                report,
                "blocked_examples",
                {
                    "field_name": header,
                    "value": raw_value,
                    "reason": describe_entity_reason(reason),
                    "source_file": source_file,
                },
            )

        if entity_headers and not any(clean_cell_value(values.get(header, "")) for header in entity_headers):
            if any(
                clean_cell_value(values.get(header, ""))
                for header in headers
                if header not in entity_headers
            ):
                report["dropped_records"] = int(report.get("dropped_records", 0)) + 1
                continue

        if _record_has_values(record):
            cleaned_records.append(record)

    result.entity_legality_report = report
    if report.get("blocked_count"):
        examples = ", ".join(
            f"{item.get('field_name')}={item.get('value')}"
            for item in report.get("blocked_examples", [])[:3]
        )
        warning = (
            f"实体合法性校验已阻断 {report.get('blocked_count', 0)} 个疑似伪实体"
            f"{f'，示例: {examples}' if examples else ''}"
        )
        result.warnings = [
            item for item in result.warnings
            if not item.startswith("实体合法性校验已阻断 ")
        ]
        result.warnings.append(warning)
    return cleaned_records


def _append_entity_example(report: dict, key: str, example: dict, limit: int = 6):
    """Append a deduplicated legality example."""
    examples = report.setdefault(key, [])
    fingerprint = (
        example.get("field_name", ""),
        example.get("value", ""),
        example.get("reason", ""),
    )
    existing = {
        (
            item.get("field_name", ""),
            item.get("value", ""),
            item.get("reason", ""),
        )
        for item in examples
    }
    if fingerprint in existing or len(examples) >= limit:
        return
    examples.append(example)


def _build_field_result(
    field_name: str,
    target_location: str,
    value: Any,
    normalized_value: Any,
    record: dict,
    table_data: dict,
) -> FilledFieldResult:
    """Create a filled field result from one extracted record."""
    evidence = record.get("field_evidence", {}).get(field_name, [])
    source_file = next((item.source_file for item in evidence if item.source_file), "")
    if not source_file:
        source_file = record.get("source_file", "")
    match_method = record.get("match_methods", {}).get(field_name, table_data.get("extraction_method", ""))
    confidence = _get_field_confidence(field_name, value, table_data, record, evidence)
    if value in (None, ""):
        missing_reason = "No data found"
    elif not evidence:
        missing_reason = "No evidence"
    else:
        missing_reason = ""
    supporting_sources = _ordered_sources(
        [
            item.source_file
            for item in evidence
            if item.source_file
        ] or [record.get("source_file", "")]
    )
    value_sources = _ordered_sources(record.get("field_value_sources", {}).get(field_name, []))
    value_record_ids = _ordered_sources(record.get("field_value_record_ids", {}).get(field_name, []))
    return FilledFieldResult(
        field_name=field_name,
        target_location=target_location,
        value=value,
        normalized_value=normalized_value,
        evidence=evidence,
        confidence=confidence,
        source_file=source_file,
        supporting_sources=supporting_sources,
        value_sources=value_sources,
        value_record_ids=value_record_ids,
        match_method=match_method,
        missing_reason=missing_reason,
    )


def _get_field_confidence(
    field_name: str,
    value: Any,
    table_data: dict,
    record: dict,
    evidence: list[CandidateEvidence],
) -> float | None:
    """Resolve confidence for one field using record-level data first."""
    if value in (None, ""):
        return None
    if not evidence:
        return None
    record_confidence = record.get("field_confidence", {}).get(field_name)
    if record_confidence is not None:
        return record_confidence
    column_confidence = table_data.get("col_confidence", {}).get(field_name)
    if column_confidence is not None:
        return column_confidence

    method = table_data.get("extraction_method", "none")
    fallback_scores = {
        "rule": 0.76,
        "llm": 0.54,
        "llm_multi": 0.57,
        "qwen": 0.54,
        "qwen_multi": 0.57,
        "qwen_narrative": 0.58,
        "none": None,
    }
    return fallback_scores.get(method, 0.60)


def _merge_table_reports(result: FilledResult, extracted_data: list[dict]):
    """Merge table-level evidence and warnings into the result."""
    seen_evidence: set[tuple[str, str, str]] = set(
        (item.source_file, item.location, item.match_reason) for item in result.evidence_report
    )
    for table_data in extracted_data:
        for warning in table_data.get("warnings", []):
            if warning not in result.warnings:
                result.warnings.append(warning)
        for evidence in table_data.get("evidence", []):
            key = (evidence.source_file, evidence.location, evidence.match_reason)
            if key not in seen_evidence:
                seen_evidence.add(key)
                result.evidence_report.append(evidence)
        _merge_filter_diagnostics_into_report(result, table_data.get("filter_diagnostics", {}))


def _merge_filter_diagnostics_into_report(result: FilledResult, diagnostics: dict[str, Any]):
    """Merge extraction-stage filter diagnostics into the result legality report."""
    if not diagnostics:
        return
    report = result.entity_legality_report or {
        "validated_fields": 0,
        "blocked_count": 0,
        "dropped_records": 0,
        "blocked_examples": [],
        "accepted_examples": [],
        "filtered_records": [],
        "filter_reason_counts": {},
        "per_source_filter_reasons": {},
        "recovered_examples": [],
    }
    for item in diagnostics.get("filtered_records", []):
        _append_unique_report_item(report, "filtered_records", item)
    for item in diagnostics.get("recovered_examples", []):
        _append_unique_report_item(report, "recovered_examples", item)
    for reason, count in (diagnostics.get("filter_reason_counts", {}) or {}).items():
        report.setdefault("filter_reason_counts", {})
        report["filter_reason_counts"][reason] = int(report["filter_reason_counts"].get(reason, 0)) + int(count)
    for source_file, reason_counts in (diagnostics.get("per_source", {}) or {}).items():
        source_bucket = report.setdefault("per_source_filter_reasons", {}).setdefault(source_file, {})
        for reason, count in reason_counts.items():
            source_bucket[reason] = int(source_bucket.get(reason, 0)) + int(count)
    result.entity_legality_report = report


def _append_unique_report_item(report: dict, key: str, item: dict[str, Any], limit: int = 12):
    """Append one deduplicated report item."""
    bucket = report.setdefault(key, [])
    fingerprint = (
        item.get("record_id", ""),
        item.get("field_name", ""),
        item.get("entity_text", item.get("value", "")),
        item.get("filter_reason", item.get("reason", "")),
    )
    existing = {
        (
            entry.get("record_id", ""),
            entry.get("field_name", ""),
            entry.get("entity_text", entry.get("value", "")),
            entry.get("filter_reason", entry.get("reason", "")),
        )
        for entry in bucket
    }
    if fingerprint in existing or len(bucket) >= limit:
        return
    bucket.append(item)


def _collect_placeholder_values(extracted_data: list[dict]) -> dict[str, str]:
    """Collect placeholder values from extracted records."""
    collected: dict[str, list[str]] = {}
    for table_data in extracted_data:
        for record in _table_records(table_data):
            for field_name, value in record.get("values", {}).items():
                if not value:
                    continue
                collected.setdefault(field_name, [])
                if value not in collected[field_name]:
                    collected[field_name].append(str(value))
    return {
        field_name: values[0] if len(values) == 1 else "；".join(values[:5])
        for field_name, values in collected.items()
    }


def _replace_placeholders_in_text(content: str, placeholder_values: dict[str, str]) -> str:
    """Replace common placeholder syntaxes in plain text."""
    updated = content
    for pattern in PLACEHOLDER_PATTERNS:
        for match in list(re.finditer(pattern, updated)):
            placeholder = match.group(1)
            value = placeholder_values.get(placeholder)
            if value:
                updated = updated.replace(match.group(0), value)
    return updated


def _placeholder_confidence(field_name: str, extracted_data: list[dict]) -> float | None:
    """Use the best confidence among records for placeholder substitution."""
    values: list[float] = []
    for table_data in extracted_data:
        for record in _table_records(table_data):
            confidence = record.get("field_confidence", {}).get(field_name)
            if confidence is not None:
                values.append(confidence)
    return round(max(values), 4) if values else None


def _placeholder_evidence(field_name: str, extracted_data: list[dict]) -> list[CandidateEvidence]:
    """Reuse the first field evidence when replacing placeholders."""
    for table_data in extracted_data:
        for record in _table_records(table_data):
            evidence = record.get("field_evidence", {}).get(field_name, [])
            if evidence:
                return evidence
    return []


def _placeholder_source_file(field_name: str, extracted_data: list[dict]) -> str:
    """Resolve placeholder source file for reporting."""
    for table_data in extracted_data:
        for record in _table_records(table_data):
            if record.get("values", {}).get(field_name):
                return record.get("source_file", "")
    return ""


def _build_markdown_table(headers: list[str], records: list[dict]) -> str:
    """Build a markdown table from records."""
    header_line = "| " + " | ".join(headers) + " |"
    separator_line = "| " + " | ".join(["---"] * len(headers)) + " |"
    data_lines = []
    for record in records:
        row = [str(record["values"].get(header, "")) for header in headers]
        data_lines.append("| " + " | ".join(row) + " |")
    return "\n".join([header_line, separator_line] + data_lines)


def _find_excel_header_row(worksheet, headers: list[str]) -> int:
    """Detect the most likely header row in an Excel sheet."""
    header_row = 1
    for row_index in range(1, min(10, worksheet.max_row + 1)):
        row_values = [clean_cell_value(worksheet.cell(row_index, column).value) for column in range(1, worksheet.max_column + 1)]
        match_count = sum(1 for header in headers if header in row_values)
        if headers and match_count >= max(1, int(len(headers) * 0.5)):
            header_row = row_index
            break
    return header_row


def _map_headers_to_sheet(worksheet, header_row: int, headers: list[str]) -> dict[int, int]:
    """Map template headers to Excel sheet columns."""
    sheet_headers = [clean_cell_value(worksheet.cell(header_row, column).value) for column in range(1, worksheet.max_column + 1)]
    mapping: dict[int, int] = {}
    for header_index, header in enumerate(headers):
        for sheet_index, sheet_header in enumerate(sheet_headers, start=1):
            if header == sheet_header or (header and sheet_header and (header in sheet_header or sheet_header in header)):
                mapping[header_index] = sheet_index
                break
    return mapping


def _map_headers_to_row(headers: list[str], row_headers: list[str]) -> dict[int, int]:
    """Map template headers to an existing row header list."""
    mapping: dict[int, int] = {}
    for header_index, header in enumerate(headers):
        for row_index, row_header in enumerate(row_headers):
            if header == row_header or (header and row_header and (header in row_header or row_header in header)):
                mapping[header_index] = row_index
                break
    return mapping


def _smart_value(value: Any):
    """Convert strings to more suitable Python values for Excel cells."""
    if value in (None, ""):
        return ""
    text = str(value).strip()
    if not text:
        return ""
    try:
        cleaned = text.replace(",", "")
        if "." not in cleaned:
            return int(cleaned)
    except (ValueError, OverflowError):
        pass
    try:
        return float(text.replace(",", ""))
    except (ValueError, OverflowError):
        pass
    return text


def _non_empty_records(records: list[dict]) -> list[dict]:
    """Keep only records that contain at least one non-empty field."""
    return [record for record in records if _record_has_values(record)]


def _ordered_sources(values: list[str]) -> list[str]:
    """Deduplicate source-like strings while preserving order."""
    ordered: list[str] = []
    seen: set[str] = set()
    for value in values:
        cleaned = clean_cell_value(str(value or ""))
        if not cleaned or cleaned in seen:
            continue
        seen.add(cleaned)
        ordered.append(cleaned)
    return ordered


def _record_has_values(record: dict) -> bool:
    """Return True when a record contains meaningful values."""
    values = record.get("values", {})
    return any(value not in (None, "", "N/A") for value in values.values())


def _metric_definitions() -> dict[str, str]:
    """Shared metric definitions used by API, logs, validation, and frontend."""
    return {
        "record_count": "本次模板各目标表/区域生成的抽取记录总数；跨表按目标位置分别计数。",
        "rows_filled": "成功写入模板且至少包含一个非空值的目标行数。",
        "expected_rows": "本次模板目标应填行数；优先使用模板可写行位，其次使用实体估计。",
        "fill_rate": "rows_filled / expected_rows * 100；若无行级目标，则回退为字段填充率。",
    }
