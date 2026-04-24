"""Natural-language document operation service."""
from __future__ import annotations

import csv
import re
import uuid
from pathlib import Path
from typing import Any

from ..core.config import OUTPUT_DIR
from ..core.exceptions import DocumentOperationError
from ..schemas.models import FileRole
from ..schemas.operation_models import DocumentOperationRequest, DocumentOperationResponse
from ..utils.text_utils import clean_cell_value
from .document_service import read_document
from .normalization_service import normalize_text
from .ollama_service import get_ollama_service
from .schema_registry_service import best_field_match


def operate_document(request: DocumentOperationRequest) -> DocumentOperationResponse:
    if not request.file_path:
        raise DocumentOperationError("file_path is required")
    if not Path(request.file_path).exists():
        raise DocumentOperationError(f"File not found: {request.file_path}")
    document = read_document(request.file_path, FileRole.SOURCE)
    intent = _detect_intent(request)
    warnings: list[str] = []

    if intent == "summarize":
        result = summarize_document(request, document, warnings)
    elif intent == "extract":
        result = extract_document_fields(request, document)
    elif intent == "find":
        result = find_in_document(request, document)
    elif intent == "replace":
        output_file = replace_placeholders(request)
        result = {"replacements": request.replacements, "output_file": output_file}
    elif intent == "export_table":
        output_file = export_tables(request, document)
        result = {"output_file": output_file, "table_count": len(document.tables)}
    elif intent == "format":
        output_file = save_formatted_text(request, document)
        result = {"output_file": output_file}
    else:
        result = {
            "text_blocks": len(document.text_blocks),
            "tables": len(document.tables),
            "raw_text_preview": document.raw_text[:2000],
        }

    output_file = result.get("output_file", "") if isinstance(result, dict) else ""
    return DocumentOperationResponse(
        status="ok",
        operation=intent,
        intent=intent,
        result=result,
        output_file=output_file,
        warnings=warnings,
    )


def summarize_document(request: DocumentOperationRequest, document=None, warnings: list[str] | None = None) -> dict[str, Any]:
    document = document or read_document(request.file_path, FileRole.SOURCE)
    text = document.raw_text or "\n".join(block.content for block in document.text_blocks)
    if request.use_llm and get_ollama_service().is_available:
        parsed, error = get_ollama_service().generate_json(
            "请对文档做结构化摘要，只输出 JSON：{\"summary\":\"...\",\"key_points\":[\"...\"],\"fields\":[\"...\"]}\n"
            f"文档内容：\n{text[:5000]}",
            "你是文档摘要助手，只输出 JSON。",
            num_predict=512,
            usage_context={"stage": "document_operation", "source_file": request.file_path},
        )
        if isinstance(parsed, dict) and not error:
            return parsed
        if warnings is not None:
            warnings.append(error or "LLM 摘要失败，已回退规则摘要")
    sentences = re.split(r"(?<=[。！？.!?])\s*", normalize_text(text))
    key_points = [sentence for sentence in sentences if len(sentence) >= 12][:8]
    return {
        "summary": " ".join(key_points[:3])[:800],
        "key_points": key_points,
        "table_count": len(document.tables),
        "text_block_count": len(document.text_blocks),
    }


def extract_document_fields(request: DocumentOperationRequest, document=None) -> dict[str, Any]:
    document = document or read_document(request.file_path, FileRole.SOURCE)
    fields = request.fields or _fields_from_instruction(request.instruction)
    extracted: dict[str, list[dict[str, Any]]] = {}
    for field in fields:
        values = []
        for table in document.tables:
            matched, score = best_field_match(field, table.headers, threshold=0.45)
            if not matched:
                continue
            col_index = table.headers.index(matched)
            for row_index, row in enumerate(table.rows[:500]):
                value = row[col_index] if col_index < len(row) else ""
                if clean_cell_value(value):
                    values.append({
                        "value": value,
                        "source_file": document.source_file,
                        "location": f"table{table.table_index}.row{row_index + 1}.col{col_index + 1}",
                        "matched_field": matched,
                        "confidence": score,
                    })
        if not values:
            values.extend(_extract_labeled_text_values(field, document.raw_text, document.source_file))
        extracted[field] = values[:100]
    return {"fields": extracted, "field_count": len(fields)}


def find_in_document(request: DocumentOperationRequest, document=None) -> dict[str, Any]:
    document = document or read_document(request.file_path, FileRole.SOURCE)
    query = request.query or request.instruction
    matches = []
    for block in document.text_blocks:
        if query and query in block.content:
            matches.append({
                "source_file": document.source_file,
                "location": f"text_block{block.block_index}",
                "snippet": block.content[:500],
            })
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            row_text = " | ".join(row)
            if query and query in row_text:
                matches.append({
                    "source_file": document.source_file,
                    "location": f"table{table.table_index}.row{row_index + 1}",
                    "snippet": row_text[:500],
                })
    return {"query": query, "matches": matches[:100], "match_count": len(matches)}


def replace_placeholders(request: DocumentOperationRequest) -> str:
    path = Path(request.file_path)
    replacements = dict(request.replacements or {})
    if not replacements:
        for key, value in re.findall(r"([A-Za-z\u4e00-\u9fa5_][\w\u4e00-\u9fa5]{0,30})\s*(?:替换为|改为|=|：|:)\s*([^\n,，；;]+)", request.instruction):
            replacements[key.strip()] = value.strip()
    output_path = OUTPUT_DIR / (request.save_as or f"{path.stem}_operated_{uuid.uuid4().hex[:8]}{path.suffix}")
    if path.suffix.lower() == ".docx":
        from docx import Document

        doc = Document(str(path))
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)
        doc.save(str(output_path))
        return str(output_path)
    text = path.read_text(encoding="utf-8", errors="ignore")
    for key, value in replacements.items():
        for pattern in (f"{{{{{key}}}}}", f"{{{key}}}", f"【{key}】", key):
            text = text.replace(pattern, value)
    output_path.write_text(text, encoding="utf-8")
    return str(output_path)


def export_tables(request: DocumentOperationRequest, document=None) -> str:
    document = document or read_document(request.file_path, FileRole.SOURCE)
    output_path = OUTPUT_DIR / (request.save_as or f"{Path(request.file_path).stem}_tables_{uuid.uuid4().hex[:8]}.csv")
    with open(output_path, "w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        for table in document.tables:
            writer.writerow([f"# table {table.table_index}", table.sheet_name])
            writer.writerow(table.headers)
            writer.writerows(table.rows)
            writer.writerow([])
    return str(output_path)


def save_formatted_text(request: DocumentOperationRequest, document=None) -> str:
    document = document or read_document(request.file_path, FileRole.SOURCE)
    output_path = OUTPUT_DIR / (request.save_as or f"{Path(request.file_path).stem}_formatted_{uuid.uuid4().hex[:8]}.txt")
    parts = [normalize_text(block.content) for block in document.text_blocks if block.content]
    for table in document.tables:
        parts.append("\t".join(table.headers))
        parts.extend("\t".join(row) for row in table.rows)
    output_path.write_text("\n\n".join(part for part in parts if part), encoding="utf-8")
    return str(output_path)


def _detect_intent(request: DocumentOperationRequest) -> str:
    explicit = (request.operation or "").strip().lower()
    if explicit:
        return explicit
    text = request.instruction or ""
    if any(token in text for token in ["摘要", "总结", "概括", "summarize", "summary"]):
        return "summarize"
    if any(token in text for token in ["提取", "抽取", "字段", "extract"]):
        return "extract"
    if any(token in text for token in ["查找", "搜索", "定位", "find", "search"]):
        return "find"
    if any(token in text for token in ["替换", "占位符", "replace"]):
        return "replace"
    if any(token in text for token in ["导出表格", "导出", "export"]):
        return "export_table"
    if any(token in text for token in ["整理", "格式", "format"]):
        return "format"
    return "inspect"


def _fields_from_instruction(instruction: str) -> list[str]:
    match = re.search(r"(?:字段|提取|抽取)[：:]\s*([^\n。]+)", instruction or "")
    if not match:
        return []
    return [item.strip() for item in re.split(r"[,，、;；]", match.group(1)) if item.strip()]


def _extract_labeled_text_values(field: str, text: str, source_file: str) -> list[dict[str, Any]]:
    values = []
    for match in re.finditer(rf"{re.escape(field)}\s*[：:]\s*([^\n，,；;。]+)", text or ""):
        values.append({
            "value": clean_cell_value(match.group(1)),
            "source_file": source_file,
            "location": f"text_pos{match.start()}",
            "matched_field": field,
            "confidence": 0.7,
        })
    return values


def _replace_in_paragraph(paragraph, replacements: dict[str, str]):
    if not paragraph.runs:
        paragraph.text = _replace_text(paragraph.text, replacements)
        return
    text = "".join(run.text for run in paragraph.runs)
    updated = _replace_text(text, replacements)
    if updated == text:
        return
    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_text(text: str, replacements: dict[str, str]) -> str:
    updated = text
    for key, value in replacements.items():
        for pattern in (f"{{{{{key}}}}}", f"{{{key}}}", f"【{key}】", key):
            updated = updated.replace(pattern, value)
    return updated
