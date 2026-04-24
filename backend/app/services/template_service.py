"""Template parsing service - supports all formats as templates."""
import re
from pathlib import Path

from ..schemas.models import (
    TemplateSchema, TemplateField, TemplateTable, TemplateSection,
    StructureType, FileRole
)
from ..core.logging import logger
from ..core.exceptions import TemplateParseError
from ..utils.text_utils import clean_cell_value
from .document_service import read_document, get_file_type


def parse_template(file_path: str) -> TemplateSchema:
    """Parse any supported file as a template."""
    ftype = get_file_type(file_path)
    path = Path(file_path)
    logger.info(f"Parsing template: {path.name} (type={ftype})")

    schema = TemplateSchema(
        template_id=path.stem,
        source_file=str(path),
        file_type=ftype,
    )

    try:
        if ftype == 'excel':
            _parse_excel_template(schema, file_path)
        elif ftype == 'word':
            _parse_word_template(schema, file_path)
        elif ftype == 'markdown':
            _parse_markdown_template(schema, file_path)
        elif ftype == 'text':
            _parse_text_template(schema, file_path)
        elif ftype == 'csv':
            _parse_csv_template(schema, file_path)
        else:
            raise TemplateParseError(f"Unsupported template type: {ftype}")
    except TemplateParseError:
        raise
    except Exception as e:
        raise TemplateParseError(f"Error parsing template {path.name}: {e}")

    # Determine structure type
    if schema.tables and not schema.sections:
        schema.structure_type = StructureType.TABULAR
    elif schema.sections and not schema.tables:
        schema.structure_type = StructureType.TEXT
    else:
        schema.structure_type = StructureType.MIXED

    logger.info(f"  -> {len(schema.tables)} tables, {len(schema.fields)} fields, "
                f"{len(schema.placeholders)} placeholders, type={schema.structure_type.value}")
    return schema


def _parse_excel_template(schema: TemplateSchema, path: str):
    import pandas as pd
    xls = pd.ExcelFile(path)

    for si, sheet_name in enumerate(xls.sheet_names):
        df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
        if df.empty:
            continue

        # Find header row
        header_row_idx = 0
        for i in range(min(5, len(df))):
            non_empty = df.iloc[i].dropna()
            if len(non_empty) >= 2:
                header_row_idx = i
                break

        headers = [clean_cell_value(v) for v in df.iloc[header_row_idx]]
        data_start = header_row_idx + 1
        total_rows = len(df) - data_start

        # Check which rows are writable (empty or have placeholders)
        writable_rows = []
        for ri in range(data_start, len(df)):
            row = df.iloc[ri]
            empty_count = sum(1 for v in row if clean_cell_value(v) == "")
            if empty_count >= len(headers) * 0.5:  # more than half empty
                writable_rows.append(ri - data_start)

        tt = TemplateTable(
            table_index=si,
            sheet_name=sheet_name,
            headers=[h for h in headers if h],
            writable_rows=writable_rows if writable_rows else list(range(total_rows)),
            writable_cols=list(range(len(headers))),
            row_count=total_rows,
            col_count=len(headers),
        )
        schema.tables.append(tt)

        for ci, h in enumerate(headers):
            if h:
                schema.fields.append(TemplateField(
                    field_name=h,
                    location=f"{sheet_name}!col{ci}",
                    field_type="text",
                ))


def _parse_word_template(schema: TemplateSchema, path: str):
    from docx import Document
    doc = Document(path)
    table_descriptions = _word_table_descriptions(doc)

    # Parse paragraphs for placeholders and sections
    placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}|\{([^}]+)\}|【([^】]+)】|__+([^_]+)__+')
    all_text = []

    for pi, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if not txt:
            continue
        all_text.append(txt)

        # Check for placeholders
        matches = placeholder_pattern.findall(txt)
        for m in matches:
            ph = next((g for g in m if g), "")
            if ph:
                schema.placeholders.append(ph)
                schema.fields.append(TemplateField(
                    field_name=ph,
                    location=f"paragraph{pi}",
                    field_type="text",
                    placeholder=ph,
                ))

        # Detect section headers
        is_heading = False
        if para.style and para.style.name:
            style = para.style.name.lower()
            if 'heading' in style:
                is_heading = True

        if is_heading or (len(txt) < 50 and (txt.endswith('：') or txt.endswith(':'))):
            schema.sections.append(TemplateSection(
                section_index=pi,
                heading=txt,
            ))

    schema.raw_text = "\n".join(all_text)

    # Parse tables
    for ti, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_data)

        if not rows_data:
            continue

        headers = rows_data[0]
        description = table_descriptions.get(ti, "")

        writable_rows = [
            row_index
            for row_index, row_data in enumerate(rows_data[1:])
            if _word_row_is_writable(row_data)
        ]
        if not writable_rows:
            writable_rows = list(range(len(rows_data) - 1))
        tt = TemplateTable(
            table_index=ti,
            headers=headers,
            writable_rows=writable_rows,
            writable_cols=list(range(len(headers))),
            row_count=len(rows_data) - 1,
            col_count=len(headers),
            description=description,
        )
        schema.tables.append(tt)

        for ci, h in enumerate(headers):
            if h:
                schema.fields.append(TemplateField(
                    field_name=h,
                    location=f"table{ti}.col{ci}",
                    field_type="text",
                ))


def _word_table_descriptions(doc) -> dict[int, str]:
    """Capture the nearest preceding paragraph text for each Word table."""
    paragraph_texts = {
        id(paragraph._element): paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.text and paragraph.text.strip()
    }
    descriptions: dict[int, str] = {}
    recent_paragraphs: list[str] = []
    table_index = 0

    for child in doc.element.body.iterchildren():
        if child.tag.endswith('}p'):
            text = paragraph_texts.get(id(child), "")
            if text:
                recent_paragraphs.append(text)
                recent_paragraphs = recent_paragraphs[-3:]
            continue
        if child.tag.endswith('}tbl'):
            descriptions[table_index] = " ".join(recent_paragraphs[-2:]).strip()
            recent_paragraphs = []
            table_index += 1

    return descriptions


def _word_row_is_writable(row_data: list[str]) -> bool:
    """Treat mostly blank or placeholder rows as writable slots."""
    cells = [clean_cell_value(cell) for cell in row_data]
    if not cells:
        return False
    empty_ratio = sum(1 for cell in cells if not cell) / max(len(cells), 1)
    has_placeholder = any(re.search(r'\{\{[^}]+\}\}|\{[^}]+\}|【[^】]+】|__+[^_]+__+', cell) for cell in cells)
    return has_placeholder or empty_ratio >= 0.4


def _parse_markdown_template(schema: TemplateSchema, path: str):
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    schema.raw_text = content

    # Find placeholders
    for pattern in [r'\{\{([^}]+)\}\}', r'\{([^}]+)\}', r'【([^】]+)】', r'___+']:
        for m in re.finditer(pattern, content):
            ph = m.group(1) if m.lastindex else "blank"
            schema.placeholders.append(ph)
            schema.fields.append(TemplateField(
                field_name=ph, location=f"md_pos{m.start()}", placeholder=ph,
            ))

    # Parse tables
    table_pattern = r'\|(.+)\|\n\|[\s:|-]+\|\n((?:\|.+\|\n?)+)'
    for ti, match in enumerate(re.finditer(table_pattern, content)):
        headers = [h.strip() for h in match.group(1).split('|') if h.strip()]
        data_lines = match.group(2).strip().split('\n')

        tt = TemplateTable(
            table_index=ti,
            headers=headers,
            writable_rows=list(range(len(data_lines))),
            writable_cols=list(range(len(headers))),
            row_count=len(data_lines),
            col_count=len(headers),
        )
        schema.tables.append(tt)

        for ci, h in enumerate(headers):
            if h:
                schema.fields.append(TemplateField(
                    field_name=h, location=f"md_table{ti}.col{ci}",
                ))

    # Parse sections
    for mi, m in enumerate(re.finditer(r'^(#{1,6})\s+(.+)', content, re.MULTILINE)):
        schema.sections.append(TemplateSection(
            section_index=mi,
            heading=m.group(2).strip(),
        ))


def _parse_text_template(schema: TemplateSchema, path: str):
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    schema.raw_text = content

    # Find placeholders
    for pattern in [r'\{\{([^}]+)\}\}', r'\{([^}]+)\}', r'【([^】]+)】',
                     r'___+\s*(\S+)?', r'<([^>]+)>']:
        for m in re.finditer(pattern, content):
            ph = m.group(1) if m.lastindex and m.group(1) else "blank"
            if ph not in schema.placeholders:
                schema.placeholders.append(ph)
                schema.fields.append(TemplateField(
                    field_name=ph, location=f"txt_pos{m.start()}", placeholder=ph,
                ))

    # Check for tabular structure
    lines = content.strip().split('\n')
    if len(lines) >= 2:
        first = lines[0]
        if '\t' in first or (first.count(',') >= 2 and '，' not in first):
            delim = '\t' if '\t' in first else ','
            headers = [h.strip() for h in first.split(delim)]
            tt = TemplateTable(
                table_index=0,
                headers=headers,
                writable_rows=list(range(len(lines) - 1)),
                writable_cols=list(range(len(headers))),
                row_count=len(lines) - 1,
                col_count=len(headers),
            )
            schema.tables.append(tt)
            for ci, h in enumerate(headers):
                if h:
                    schema.fields.append(TemplateField(
                        field_name=h, location=f"txt_table0.col{ci}",
                    ))

    # Sections by double newline
    blocks = re.split(r'\n\s*\n', content)
    for bi, block in enumerate(blocks):
        block = block.strip()
        if block:
            schema.sections.append(TemplateSection(
                section_index=bi, heading=block.split('\n')[0][:80],
                content_template=block,
            ))


def _parse_csv_template(schema: TemplateSchema, path: str):
    import csv as csv_mod
    with open(path, 'r', encoding='utf-8') as f:
        reader = csv_mod.reader(f)
        rows = list(reader)

    if not rows:
        return

    headers = [c.strip() for c in rows[0]]
    tt = TemplateTable(
        table_index=0,
        headers=headers,
        writable_rows=list(range(len(rows) - 1)),
        writable_cols=list(range(len(headers))),
        row_count=len(rows) - 1,
        col_count=len(headers),
    )
    schema.tables.append(tt)

    for ci, h in enumerate(headers):
        if h:
            schema.fields.append(TemplateField(
                field_name=h, location=f"csv_col{ci}",
            ))
    schema.raw_text = "\n".join([",".join(r) for r in rows])
