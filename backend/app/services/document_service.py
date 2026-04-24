"""Unified document reading and normalization service."""
import os
import re
import csv
from pathlib import Path
from typing import Optional

import pandas as pd

from ..schemas.models import DocumentBundle, TextBlock, NormalizedTable, FileRole
from ..core.logging import logger
from ..core.exceptions import DocumentReadError
from ..utils.text_utils import clean_cell_value


# Supported file extensions
SUPPORTED_EXTENSIONS = {'.xlsx', '.xls', '.docx', '.md', '.txt', '.csv', '.json'}


def get_file_type(path: str) -> str:
    ext = Path(path).suffix.lower()
    mapping = {
        '.xlsx': 'excel', '.xls': 'excel',
        '.docx': 'word',
        '.md': 'markdown',
        '.txt': 'text',
        '.csv': 'csv',
        '.json': 'json',
    }
    return mapping.get(ext, 'unknown')


def read_document(file_path: str, role: FileRole | str = FileRole.UNKNOWN) -> DocumentBundle:
    """Read any supported document and normalize it into DocumentBundle."""
    if isinstance(role, str):
        valid_roles = {member.value for member in FileRole}
        role = FileRole(role) if role in valid_roles else FileRole.UNKNOWN
    path = Path(file_path)
    if not path.exists():
        raise DocumentReadError(f"File not found: {file_path}")

    ftype = get_file_type(file_path)
    doc_id = path.stem

    logger.info(f"Reading document: {path.name} (type={ftype}, role={role.value})")

    bundle = DocumentBundle(
        document_id=doc_id,
        source_file=str(path),
        file_type=ftype,
        role=role,
    )

    try:
        if ftype == 'excel':
            _read_excel(bundle, file_path)
        elif ftype == 'word':
            _read_word(bundle, file_path)
        elif ftype == 'markdown':
            _read_markdown(bundle, file_path)
        elif ftype == 'text':
            _read_text(bundle, file_path)
        elif ftype == 'csv':
            _read_csv(bundle, file_path)
        elif ftype == 'json':
            _read_json(bundle, file_path)
        else:
            raise DocumentReadError(f"Unsupported file type: {ftype}")
    except DocumentReadError:
        raise
    except Exception as e:
        raise DocumentReadError(f"Error reading {path.name}: {e}")

    logger.info(f"  -> {len(bundle.text_blocks)} text blocks, {len(bundle.tables)} tables")
    bundle.metadata.setdefault("title", _guess_title(bundle))
    bundle.metadata.setdefault("inferred_topic", _infer_document_topic(bundle))
    bundle.metadata.setdefault("topic_tokens", sorted(_topic_tokens(bundle.metadata.get("inferred_topic", ""))))
    bundle.metadata.setdefault("role", bundle.role.value)
    bundle.metadata.setdefault("text_block_count", len(bundle.text_blocks))
    bundle.metadata.setdefault("table_count", len(bundle.tables))
    return bundle


def _read_excel(bundle: DocumentBundle, path: str):
    xls = pd.ExcelFile(path)
    all_text_parts = []

    for si, sheet_name in enumerate(xls.sheet_names):
        df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
        if df.empty:
            continue

        # Try to find header row (first non-empty row)
        header_row_idx = 0
        for i in range(min(5, len(df))):
            non_empty = df.iloc[i].dropna()
            if len(non_empty) >= 2:
                header_row_idx = i
                break

        headers = [clean_cell_value(v) for v in df.iloc[header_row_idx]]
        data_rows = []
        for ri in range(header_row_idx + 1, len(df)):
            row = [clean_cell_value(v) for v in df.iloc[ri]]
            if any(c for c in row):
                data_rows.append(row)

        table = NormalizedTable(
            table_index=si,
            sheet_name=sheet_name,
            headers=headers,
            rows=data_rows,
            row_count=len(data_rows),
            col_count=len(headers),
        )
        bundle.tables.append(table)

        # Build text representation
        text_repr = f"[Sheet: {sheet_name}]\n"
        text_repr += " | ".join(headers) + "\n"
        for r in data_rows[:100]:  # limit rows for text
            text_repr += " | ".join(r) + "\n"
        all_text_parts.append(text_repr)
        bundle.text_blocks.append(TextBlock(
            content=text_repr.strip(),
            heading_level=1,
            block_index=len(bundle.text_blocks),
        ))

    bundle.raw_text = "\n\n".join(all_text_parts)
    if not bundle.text_blocks and bundle.raw_text:
        bundle.text_blocks = [TextBlock(content=bundle.raw_text, block_index=0)]


def _read_word(bundle: DocumentBundle, path: str):
    from docx import Document
    doc = Document(path)

    text_parts = []
    for pi, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if not txt:
            continue
        # Detect heading level
        level = 0
        if para.style and para.style.name:
            style = para.style.name.lower()
            if 'heading' in style:
                for ch in style:
                    if ch.isdigit():
                        level = int(ch)
                        break
                if level == 0:
                    level = 1

        bundle.text_blocks.append(TextBlock(
            content=txt, heading_level=level, block_index=pi
        ))
        text_parts.append(txt)

    # Read tables
    for ti, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_data)

        if not rows_data:
            continue

        headers = rows_data[0]
        data_rows = rows_data[1:]

        nt = NormalizedTable(
            table_index=ti,
            headers=headers,
            rows=data_rows,
            row_count=len(data_rows),
            col_count=len(headers),
        )
        bundle.tables.append(nt)

        text_repr = f"[Word Table {ti}]\n"
        text_repr += " | ".join(headers) + "\n"
        for r in data_rows:
            text_repr += " | ".join(r) + "\n"
        text_parts.append(text_repr)

    bundle.raw_text = "\n".join(text_parts)
    bundle.metadata["title"] = text_parts[0] if text_parts else Path(path).stem


def _read_markdown(bundle: DocumentBundle, path: str):
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()

    bundle.raw_text = content

    # Parse headings and text blocks
    lines = content.split('\n')
    current_block = []
    block_idx = 0

    for line in lines:
        heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if heading_match:
            if current_block:
                bundle.text_blocks.append(TextBlock(
                    content='\n'.join(current_block),
                    heading_level=0,
                    block_index=block_idx,
                ))
                block_idx += 1
                current_block = []

            level = len(heading_match.group(1))
            bundle.text_blocks.append(TextBlock(
                content=heading_match.group(2).strip(),
                heading_level=level,
                block_index=block_idx,
            ))
            block_idx += 1
        else:
            if line.strip():
                current_block.append(line)

    if current_block:
        bundle.text_blocks.append(TextBlock(
            content='\n'.join(current_block),
            heading_level=0,
            block_index=block_idx,
        ))

    # Parse markdown tables
    table_pattern = r'\|(.+)\|\n\|[\s:|-]+\|\n((?:\|.+\|\n?)+)'
    for ti, match in enumerate(re.finditer(table_pattern, content)):
        headers = [h.strip() for h in match.group(1).split('|') if h.strip()]
        data_rows = []
        for line in match.group(2).strip().split('\n'):
            row = [c.strip() for c in line.split('|') if c.strip()]
            if row:
                data_rows.append(row)

        nt = NormalizedTable(
            table_index=ti,
            headers=headers,
            rows=data_rows,
            row_count=len(data_rows),
            col_count=len(headers),
        )
        bundle.tables.append(nt)


def _read_text(bundle: DocumentBundle, path: str):
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()

    bundle.raw_text = content

    # Check if it's structured (tab/comma separated)
    lines = content.strip().split('\n')
    if len(lines) >= 2:
        first_line = lines[0]
        if '\t' in first_line or (first_line.count(',') >= 2 and '，' not in first_line):
            delim = '\t' if '\t' in first_line else ','
            headers = [h.strip() for h in first_line.split(delim)]
            data_rows = []
            for line in lines[1:]:
                if line.strip():
                    row = [c.strip() for c in line.split(delim)]
                    data_rows.append(row)
            if headers and data_rows:
                bundle.tables.append(NormalizedTable(
                    table_index=0,
                    headers=headers,
                    rows=data_rows,
                    row_count=len(data_rows),
                    col_count=len(headers),
                ))

    # Split into text blocks by paragraph
    paragraphs = re.split(r'\n\s*\n', content)
    non_empty_lines = [line.strip() for line in content.splitlines() if line.strip()]
    if len([para for para in paragraphs if para.strip()]) <= 1 and len(non_empty_lines) >= 2:
        paragraphs = non_empty_lines
    for bi, para in enumerate(paragraphs):
        para = para.strip()
        if para:
            bundle.text_blocks.append(TextBlock(
                content=para, heading_level=0, block_index=bi
            ))
    bundle.metadata["title"] = bundle.text_blocks[0].content.splitlines()[0][:80] if bundle.text_blocks else Path(path).stem


def _read_csv(bundle: DocumentBundle, path: str):
    with open(path, 'r', encoding='utf-8') as f:
        # Detect dialect
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            dialect = csv.excel

        reader = csv.reader(f, dialect)
        rows = list(reader)

    if not rows:
        bundle.raw_text = ""
        return

    headers = [c.strip() for c in rows[0]]
    data_rows = [[c.strip() for c in r] for r in rows[1:] if any(c.strip() for c in r)]

    bundle.tables.append(NormalizedTable(
        table_index=0,
        headers=headers,
        rows=data_rows,
        row_count=len(data_rows),
        col_count=len(headers),
    ))

    text_parts = [" | ".join(headers)]
    for r in data_rows:
        text_parts.append(" | ".join(r))
    bundle.raw_text = "\n".join(text_parts)
    bundle.text_blocks = [TextBlock(content=bundle.raw_text, block_index=0)]
    bundle.metadata["title"] = headers[0] if headers else Path(path).stem


def _read_json(bundle: DocumentBundle, path: str):
    import json

    try:
        with open(path, "r", encoding="utf-8") as handle:
            payload = json.load(handle)
    except UnicodeDecodeError:
        with open(path, "r", encoding="utf-8", errors="ignore") as handle:
            payload = json.load(handle)

    _populate_bundle_from_json(bundle, payload, title=Path(path).stem)


def _populate_bundle_from_json(bundle: DocumentBundle, payload, title: str = "json"):
    """Convert arbitrary JSON payloads into text and table channels."""
    rows = _json_records(payload)
    if rows:
        headers = _ordered_json_headers(rows)
        table_rows = [
            [clean_cell_value(_json_scalar(row.get(header, ""))) for header in headers]
            for row in rows
        ]
        bundle.tables.append(NormalizedTable(
            table_index=0,
            sheet_name=title,
            headers=headers,
            rows=table_rows,
            row_count=len(table_rows),
            col_count=len(headers),
        ))
        text_lines = [" | ".join(headers)]
        text_lines.extend(" | ".join(row) for row in table_rows[:200])
        bundle.raw_text = "\n".join(text_lines)
    else:
        import json

        bundle.raw_text = json.dumps(payload, ensure_ascii=False, indent=2)

    if bundle.raw_text:
        bundle.text_blocks.append(TextBlock(
            content=bundle.raw_text[:20000],
            heading_level=0,
            block_index=0,
        ))
    bundle.metadata["title"] = title
    bundle.metadata["json_root_type"] = type(payload).__name__


def _json_records(payload) -> list[dict]:
    if isinstance(payload, list):
        if all(isinstance(item, dict) for item in payload):
            return [_flatten_json_object(item) for item in payload]
        return [{"value": _json_scalar(item)} for item in payload]
    if isinstance(payload, dict):
        for value in payload.values():
            if isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
                return [_flatten_json_object(item) for item in value]
        return [_flatten_json_object(payload)]
    return []


def _flatten_json_object(obj: dict, prefix: str = "", max_depth: int = 3) -> dict[str, str]:
    flattened: dict[str, str] = {}
    for key, value in obj.items():
        current_key = f"{prefix}.{key}" if prefix else str(key)
        if isinstance(value, dict) and max_depth > 0:
            flattened.update(_flatten_json_object(value, current_key, max_depth=max_depth - 1))
        elif isinstance(value, list):
            flattened[current_key] = "；".join(clean_cell_value(_json_scalar(item)) for item in value[:20])
        else:
            flattened[current_key] = clean_cell_value(_json_scalar(value))
    return flattened


def _json_scalar(value) -> str:
    if isinstance(value, (dict, list)):
        import json

        return json.dumps(value, ensure_ascii=False)
    return "" if value is None else str(value)


def _ordered_json_headers(rows: list[dict]) -> list[str]:
    headers: list[str] = []
    seen: set[str] = set()
    for row in rows:
        for key in row.keys():
            if key in seen:
                continue
            seen.add(key)
            headers.append(key)
    return headers


def _guess_title(bundle: DocumentBundle) -> str:
    """Return a short document title for diagnostics and mismatch checks."""
    for block in bundle.text_blocks[:6]:
        line = block.content.strip().splitlines()[0].strip()
        if 2 <= len(line) <= 80:
            return line
    return Path(bundle.source_file).stem


def _infer_document_topic(bundle: DocumentBundle) -> str:
    """Infer a stable topic string from title and early content."""
    candidates = [
        str(bundle.metadata.get("title", "")),
        Path(bundle.source_file).stem,
    ]
    for block in bundle.text_blocks[:4]:
        line = clean_cell_value(block.content.splitlines()[0] if block.content else "")
        if line:
            candidates.append(line)

    for candidate in candidates:
        normalized = clean_cell_value(candidate)
        if 2 <= len(normalized) <= 80:
            return normalized
    return Path(bundle.source_file).stem


def _topic_tokens(text: str) -> set[str]:
    """Extract lightweight topic tokens for later source-template matching."""
    tokens = set(re.findall(r'[A-Za-z]{2,20}|[\u4e00-\u9fa5]{2,16}', text or ""))
    return {
        token.lower()
        for token in tokens
        if len(token.strip()) >= 2
    }
