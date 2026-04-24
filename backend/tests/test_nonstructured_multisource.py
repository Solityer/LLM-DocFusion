"""Tests for non-structured file recognition, multi-source joint filling, and fill correctness."""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import tempfile
from pathlib import Path


# ---------------------------------------------------------------------------
# 1. Non-structured files (txt/md) recognized as narrative sources
# ---------------------------------------------------------------------------
def test_txt_and_md_detected_as_narrative_source():
    """txt and md files should be detected as narrative sources."""
    from app.schemas.models import DocumentBundle, TextBlock
    from app.services.extraction_service import _is_narrative_source

    txt_doc = DocumentBundle(
        file_type="text",
        text_blocks=[TextBlock(content="2024年GDP为1200亿元。", block_index=0)],
        tables=[],
    )
    assert _is_narrative_source(txt_doc) is True

    md_doc = DocumentBundle(
        file_type="markdown",
        text_blocks=[TextBlock(content="# 经济报告\n全年GDP达到5000亿。", block_index=0)],
        tables=[],
    )
    assert _is_narrative_source(md_doc) is True

    # A document with many text blocks and at most one table should also be narrative
    mixed_doc = DocumentBundle(
        file_type="text",
        text_blocks=[TextBlock(content=f"段落{i}，数据{i*100}万。", block_index=i) for i in range(6)],
        tables=[],
    )
    assert _is_narrative_source(mixed_doc) is True

    print("✓ test_txt_and_md_detected_as_narrative_source passed")


# ---------------------------------------------------------------------------
# 2. Multi-source records get merged by entity key
# ---------------------------------------------------------------------------
def test_multisource_records_merge_by_entity_key():
    """Records from different source files with the same entity key should merge."""
    from app.schemas.models import CandidateEvidence, RequirementSpec
    from app.services.extraction_service import _merge_records_by_semantic_key

    headers = ["城市", "GDP（亿元）", "人口（万人）"]
    record_a = {
        "values": {"城市": "北京", "GDP（亿元）": "41000", "人口（万人）": ""},
        "field_confidence": {"城市": 0.9, "GDP（亿元）": 0.85, "人口（万人）": None},
        "field_evidence": {
            "城市": [CandidateEvidence(source_file="/tmp/src-a.docx", location="t0", raw_snippet="北京GDP41000亿", match_reason="rule", confidence=0.9)],
            "GDP（亿元）": [CandidateEvidence(source_file="/tmp/src-a.docx", location="t0", raw_snippet="北京GDP41000亿", match_reason="rule", confidence=0.85)],
        },
        "match_methods": {"城市": "text_exact", "GDP（亿元）": "text_exact", "人口（万人）": "none"},
        "source_file": "/tmp/src-a.docx",
        "source_location": "text_block0",
        "row_index": 0,
    }
    record_b = {
        "values": {"城市": "北京", "GDP（亿元）": "", "人口（万人）": "2189"},
        "field_confidence": {"城市": 0.9, "GDP（亿元）": None, "人口（万人）": 0.8},
        "field_evidence": {
            "城市": [CandidateEvidence(source_file="/tmp/src-b.md", location="t0", raw_snippet="北京人口2189万", match_reason="rule", confidence=0.9)],
            "人口（万人）": [CandidateEvidence(source_file="/tmp/src-b.md", location="t0", raw_snippet="北京人口2189万", match_reason="rule", confidence=0.8)],
        },
        "match_methods": {"城市": "text_exact", "GDP（亿元）": "none", "人口（万人）": "text_exact"},
        "source_file": "/tmp/src-b.md",
        "source_location": "text_block0",
        "row_index": 0,
    }
    requirement = RequirementSpec(raw_text="帮我填表")
    merged = _merge_records_by_semantic_key([record_a, record_b], headers, requirement, use_llm=False)

    # Should merge into a single record since both have 城市=北京
    assert len(merged) == 1, f"Expected 1 merged record but got {len(merged)}"
    values = merged[0]["values"]
    assert values["城市"] == "北京"
    assert values["GDP（亿元）"] == "41000"
    assert values["人口（万人）"] == "2189"

    print("✓ test_multisource_records_merge_by_entity_key passed")


# ---------------------------------------------------------------------------
# 3. Word template auto-expansion when records exceed writable rows
# ---------------------------------------------------------------------------
def test_word_fill_auto_expands_rows():
    """Word template with fewer writable rows than records should auto-expand."""
    from docx import Document

    from app.schemas.models import TemplateSchema, TemplateTable
    from app.services.fill_service import fill_template

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        template_path = temp_root / "template.docx"
        output_path = temp_root / "output.docx"

        # Create a Word template with 1 header row and 2 data rows
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "国家"
        table.rows[0].cells[1].text = "确诊数"
        # Row 1 and 2 are writable (empty)
        doc.save(str(template_path))

        template = TemplateSchema(
            source_file=str(template_path),
            file_type="word",
            tables=[
                TemplateTable(
                    table_index=0,
                    headers=["国家", "确诊数"],
                    writable_rows=[0, 1],  # Only 2 writable rows
                    writable_cols=[0, 1],
                    row_count=2,
                    col_count=2,
                )
            ],
        )

        # 4 records, more than the 2 writable rows
        records = [
            {"values": {"国家": "美国", "确诊数": "100000"}, "field_confidence": {}, "field_evidence": {}, "match_methods": {}, "source_file": "/tmp/s.txt"},
            {"values": {"国家": "巴西", "确诊数": "80000"}, "field_confidence": {}, "field_evidence": {}, "match_methods": {}, "source_file": "/tmp/s.txt"},
            {"values": {"国家": "印度", "确诊数": "70000"}, "field_confidence": {}, "field_evidence": {}, "match_methods": {}, "source_file": "/tmp/s.txt"},
            {"values": {"国家": "俄罗斯", "确诊数": "60000"}, "field_confidence": {}, "field_evidence": {}, "match_methods": {}, "source_file": "/tmp/s.txt"},
        ]
        extracted = [{
            "table_index": 0,
            "headers": ["国家", "确诊数"],
            "records": records,
            "col_confidence": {},
            "extraction_method": "rule",
        }]

        result = fill_template(template, extracted, str(output_path))

        # Verify all 4 records were written
        assert result.rows_filled == 4
        assert result.record_count == 4

        # Verify the output document has the auto-expanded rows
        output_doc = Document(str(output_path))
        output_table = output_doc.tables[0]
        assert len(output_table.rows) >= 5  # 1 header + 4 data rows

        # Verify content
        assert output_table.rows[1].cells[0].text == "美国"
        assert output_table.rows[2].cells[0].text == "巴西"
        # Rows 3 and 4 are auto-expanded
        assert output_table.rows[3].cells[0].text == "印度"
        assert output_table.rows[4].cells[0].text == "俄罗斯"

        # Should have a warning about auto-expansion
        assert any("自动扩展" in w for w in result.warnings)

    print("✓ test_word_fill_auto_expands_rows passed")


# ---------------------------------------------------------------------------
# 4. TXT fill preserves non-table text
# ---------------------------------------------------------------------------
def test_txt_fill_preserves_surrounding_text():
    """Filling a TXT template should only replace the table region, not the surrounding text."""
    from app.schemas.models import TemplateSchema, TemplateTable
    from app.services.fill_service import fill_template

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        template_path = temp_root / "template.txt"
        output_path = temp_root / "output.txt"

        # Create a TXT template with text before and after the table
        template_content = (
            "这是报告标题\n"
            "日期：2024年\n"
            "\n"
            "国家\t确诊数\t死亡数\n"
            "美国\t\t\n"
            "巴西\t\t\n"
            "\n"
            "数据来源：WHO统计报告\n"
            "备注：以上数据仅供参考\n"
        )
        template_path.write_text(template_content, encoding="utf-8")

        template = TemplateSchema(
            source_file=str(template_path),
            file_type="text",
            raw_text=template_content,
            tables=[
                TemplateTable(
                    table_index=0,
                    headers=["国家", "确诊数", "死亡数"],
                    writable_rows=[0, 1],
                    writable_cols=[0, 1, 2],
                    row_count=2,
                    col_count=3,
                )
            ],
        )

        records = [
            {"values": {"国家": "美国", "确诊数": "100000", "死亡数": "3000"}, "field_confidence": {}, "field_evidence": {}, "match_methods": {}, "source_file": "/tmp/s.txt"},
            {"values": {"国家": "巴西", "确诊数": "80000", "死亡数": "2500"}, "field_confidence": {}, "field_evidence": {}, "match_methods": {}, "source_file": "/tmp/s.txt"},
        ]
        extracted = [{
            "table_index": 0,
            "headers": ["国家", "确诊数", "死亡数"],
            "records": records,
            "col_confidence": {},
            "extraction_method": "rule",
        }]

        result = fill_template(template, extracted, str(output_path))

        output_content = output_path.read_text(encoding="utf-8")

        # Non-table text must be preserved
        assert "这是报告标题" in output_content, "Report title lost!"
        assert "日期：2024年" in output_content, "Date line lost!"
        assert "数据来源：WHO统计报告" in output_content, "Data source line lost!"
        assert "备注：以上数据仅供参考" in output_content, "Remark line lost!"

        # Table data should be filled
        assert "100000" in output_content
        assert "80000" in output_content
        assert "3000" in output_content

    print("✓ test_txt_fill_preserves_surrounding_text passed")


# ---------------------------------------------------------------------------
# 5. Relaxed segment relevance still accepts valid entity+number segments
# ---------------------------------------------------------------------------
def test_segment_relevance_accepts_entity_with_number():
    """A text segment with at least one alias hit plus entity and number should be relevant."""
    from app.schemas.models import RequirementSpec
    from app.services.extraction_service import _segment_is_relevant

    # A segment with entity and numbers that has an alias hit
    result = _segment_is_relevant(
        "北京市2024年GDP总量为41610亿元，同比增长5.2%。",
        headers=["城市", "GDP总量（亿元）", "增长率"],
        requirement=RequirementSpec(raw_text="帮我填表"),
        template_context={"topic_tokens": {"GDP", "经济", "增长"}},
    )
    assert result is True, "Segment with alias hit, entity, and number should be relevant"

    print("✓ test_segment_relevance_accepts_entity_with_number passed")


# ---------------------------------------------------------------------------
# 6. Source context scoring accepts matching narrative sources
# ---------------------------------------------------------------------------
def test_source_context_accepts_relevant_narrative():
    """A narrative source whose content overlaps with template indicators should be accepted."""
    from app.schemas.models import DocumentBundle, RequirementSpec, TextBlock
    from app.services.extraction_service import _source_context_score

    doc = DocumentBundle(
        source_file="/tmp/health_report.txt",
        file_type="text",
        text_blocks=[
            TextBlock(content="2024年卫生健康事业发展统计公报", heading_level=1, block_index=0),
            TextBlock(content="全国医疗卫生机构总诊疗人次达95.6亿人次，比上年增加4.5亿人次，增长4.9%。", block_index=1),
            TextBlock(content="全国卫生人员总数达1246.3万人，其中卫生技术人员1074.2万人。", block_index=2),
            TextBlock(content="医院34523个，其中公立医院11724个。", block_index=3),
        ],
        tables=[],
    )

    score = _source_context_score(
        doc,
        headers=["诊疗人次", "卫生人员", "医院数"],
        requirement=RequirementSpec(raw_text="帮我填表", entity_keywords=["全国"]),
        template_context={"topic_tokens": {"卫生", "健康", "统计"}},
    )
    assert score >= 3, f"Score {score} too low for clearly matching narrative source"

    print("✓ test_source_context_accepts_relevant_narrative passed")


# ---------------------------------------------------------------------------
# 7. Cross-topic sources still blocked after relaxation
# ---------------------------------------------------------------------------
def test_relaxed_filtering_blocks_cross_topic_source():
    """Even with relaxed thresholds, a completely off-topic source should be rejected."""
    from app.schemas.models import DocumentBundle, RequirementSpec, TextBlock
    from app.services.extraction_service import _source_context_score

    # A sports document should not match a COVID template context
    sports_doc = DocumentBundle(
        source_file="/tmp/sports.txt",
        file_type="text",
        text_blocks=[
            TextBlock(content="2024年中国体育产业发展报告", heading_level=1, block_index=0),
            TextBlock(content="全年体育彩票销售额达到2345亿元。", block_index=1),
            TextBlock(content="运动员注册总数为120万人。", block_index=2),
        ],
        tables=[],
    )

    score = _source_context_score(
        sports_doc,
        headers=["确诊病例", "死亡病例", "治愈数", "检测数"],
        requirement=RequirementSpec(raw_text="填写COVID数据", entity_keywords=["国家", "地区"]),
        template_context={"topic_tokens": {"COVID", "疫情", "新冠"}},
    )
    assert score < 3, f"Cross-topic sports doc scored {score}, should be below threshold"

    print("✓ test_relaxed_filtering_blocks_cross_topic_source passed")


if __name__ == "__main__":
    test_txt_and_md_detected_as_narrative_source()
    test_multisource_records_merge_by_entity_key()
    test_word_fill_auto_expands_rows()
    test_txt_fill_preserves_surrounding_text()
    test_segment_relevance_accepts_entity_with_number()
    test_source_context_accepts_relevant_narrative()
    test_relaxed_filtering_blocks_cross_topic_source()
    print("\n✅ All non-structured & multi-source tests passed!")
