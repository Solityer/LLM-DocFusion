#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
通用数据填充框架 - 支持Excel、Word、Markdown、Text四种文件类型全组合填充
数据源类型: Excel(.xlsx/.xls)、Word(.docx)、Markdown(.md)、Text(.txt)
模板类型: Excel(.xlsx/.xls)、Word(.docx)
支持所有8种组合的数据填充
"""

import os
import re
import sys
import json
import argparse
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from difflib import SequenceMatcher
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple


class UniversalDataFiller:
    
    DATA_FORMATS = {
        'excel': ['.xlsx', '.xls'],
        'word': ['.docx'],
        'markdown': ['.md'],
        'text': ['.txt']
    }
    
    TEMPLATE_FORMATS = {
        'excel': ['.xlsx', '.xls'],
        'word': ['.docx']
    }
    
    def __init__(self, config: Optional[Dict] = None):
        self.config = config or {}
    
    def get_file_type(self, file_path: str) -> str:
        """获取文件类型"""
        ext = Path(file_path).suffix.lower()
        for ftype, extensions in self.DATA_FORMATS.items():
            if ext in extensions:
                return ftype
        for ftype, extensions in self.TEMPLATE_FORMATS.items():
            if ext in extensions:
                return ftype
        return 'unknown'
    
    def similarity(self, str1: str, str2: str) -> float:
        """计算字符串相似度"""
        return SequenceMatcher(None, str(str1), str(str2)).ratio()
    
    def find_best_match(self, target: str, candidates: List[str], threshold: float = 0.6) -> Tuple[Optional[str], float]:
        """找到最佳匹配"""
        best_match = None
        best_score = 0
        
        for candidate in candidates:
            score = self.similarity(target, candidate)
            if score > best_score and score >= threshold:
                best_score = score
                best_match = candidate
        
        return best_match, best_score
    
    def extract_numbers_from_text(self, text: str) -> List[str]:
        """从文本中提取所有数字"""
        number_pattern = r'[0-9,\.]+'
        numbers = re.findall(number_pattern, text)
        return [n.replace(',', '') for n in numbers if n.replace(',', '').replace('.', '').isdigit()]
    
    def extract_dates_from_text(self, text: str) -> List[str]:
        """从文本中提取所有日期（可选功能，按需调用）"""
        date_patterns = [
            r'\d{4}[年/-]\d{1,2}[月/-]\d{1,2}[日号]?',
            r'\d{4}[年/-]\d{1,2}[月/-]',
            r'\d{1,2}[月/-]\d{1,2}[日号]?',
        ]
        dates = []
        for pattern in date_patterns:
            dates.extend(re.findall(pattern, text))
        return dates
    
    def parse_requirement(self, requirement: str) -> Dict[str, Any]:
        """解析用户要求"""
        params = {
            'date_range': None,
            'filters': {},
            'keywords': []
        }
        
        if not requirement:
            return params
        
        date_range = self._extract_date_range(requirement)
        if date_range:
            params['date_range'] = date_range
        
        filters = self._extract_filters(requirement)
        if filters:
            params['filters'] = filters
        
        keywords = self._extract_keywords(requirement)
        if keywords:
            params['keywords'] = keywords
        
        return params
    
    def _extract_date_range(self, text: str) -> Optional[Tuple]:
        """提取日期范围"""
        range_pattern = r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})?[日号]?\s*[到至~－-]\s*(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})?[日号]?'
        match = re.search(range_pattern, text)
        
        if match:
            start = f"{match.group(1)}-{match.group(2).zfill(2)}-{match.group(3).zfill(2) if match.group(3) else '01'}"
            end = f"{match.group(4)}-{match.group(5).zfill(2)}-{match.group(6).zfill(2) if match.group(6) else '01'}"
            return (pd.to_datetime(start), pd.to_datetime(end))
        
        single_pattern = r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})?[日号]?'
        match = re.search(single_pattern, text)
        if match:
            date = f"{match.group(1)}-{match.group(2).zfill(2)}-{match.group(3).zfill(2) if match.group(3) else '01'}"
            dt = pd.to_datetime(date)
            return (dt, dt)
        
        return None
    
    def _extract_filters(self, text: str) -> Dict[str, List[str]]:
        """提取过滤条件"""
        filters = {}
        
        table_city_pattern = r'表[一二三四五六七八九十]+[：:][^城]*城市[：:]\s*([^\s，。；\n]+)'
        table_matches = re.findall(table_city_pattern, text)
        if table_matches:
            filters['城市'] = [c.strip() for c in table_matches if c.strip()]
            return filters
        
        simple_city_pattern = r'城市[：:]\s*([^\s，。；\n]+)'
        simple_matches = re.findall(simple_city_pattern, text)
        if simple_matches:
            filters['城市'] = [c.strip() for c in simple_matches if c.strip()]
            return filters
        
        return filters
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        keywords = []
        keyword_pattern = r'(?:提取|分析|统计|汇总)[：:]?\s*([^\n，。]+)'
        match = re.search(keyword_pattern, text)
        if match:
            keywords = [k.strip() for k in match.group(1).split() if k.strip()]
        return keywords
    
    def read_excel_data(self, file_path: str) -> Dict[str, Any]:
        """读取Excel数据源"""
        result = {
            'path': file_path,
            'type': 'excel',
            'data': None,
            'columns': [],
            'text': None
        }
        
        try:
            df = pd.read_excel(file_path)
            result['data'] = df
            result['columns'] = list(df.columns)
            result['text'] = df.to_string()
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_word_data(self, file_path: str) -> Dict[str, Any]:
        """读取Word数据源"""
        result = {
            'path': file_path,
            'type': 'word',
            'data': None,
            'columns': [],
            'text': None,
            'tables': [],
            'paragraphs': []
        }
        
        try:
            doc = Document(file_path)
            result['paragraphs'] = [p.text for p in doc.paragraphs if p.text.strip()]
            result['text'] = '\n'.join(result['paragraphs'])
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                result['tables'].append(table_data)
            
            if result['tables']:
                first_table = result['tables'][0]
                if first_table:
                    result['columns'] = first_table[0]
                    result['data'] = pd.DataFrame(first_table[1:], columns=first_table[0])
                    
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_markdown_data(self, file_path: str) -> Dict[str, Any]:
        """读取Markdown数据源"""
        result = {
            'path': file_path,
            'type': 'markdown',
            'data': None,
            'columns': [],
            'text': None,
            'tables': []
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            result['text'] = content
            
            table_pattern = r'\|(.+)\|\n\|[-\s|:]+\|\n((?:\|.+\|\n?)+)'
            matches = re.findall(table_pattern, content)
            
            for match in matches:
                headers = [h.strip() for h in match[0].split('|') if h.strip()]
                rows = []
                for line in match[1].strip().split('\n'):
                    row = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if row:
                        rows.append(row)
                
                if headers and rows:
                    result['tables'].append({'headers': headers, 'rows': rows})
            
            if result['tables']:
                first_table = result['tables'][0]
                result['columns'] = first_table['headers']
                result['data'] = pd.DataFrame(first_table['rows'], columns=first_table['headers'])
                
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_text_data(self, file_path: str) -> Dict[str, Any]:
        """读取纯文本数据源"""
        result = {
            'path': file_path,
            'type': 'text',
            'data': None,
            'columns': [],
            'text': None
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            result['text'] = content
            
            lines = content.strip().split('\n')
            if lines:
                first_line = lines[0]
                if ',' in first_line or '\t' in first_line:
                    delimiter = ',' if ',' in first_line else '\t'
                    headers = [h.strip() for h in first_line.split(delimiter)]
                    rows = []
                    for line in lines[1:]:
                        if line.strip():
                            row = [cell.strip() for cell in line.split(delimiter)]
                            rows.append(row)
                    
                    if headers and rows:
                        result['columns'] = headers
                        result['data'] = pd.DataFrame(rows, columns=headers)
                        
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_data_source(self, file_path: str) -> Dict[str, Any]:
        """读取数据源 - 自动识别文件类型"""
        file_type = self.get_file_type(file_path)
        
        print(f"读取数据源: {os.path.basename(file_path)}")
        print(f"  文件类型: {file_type}")
        
        if file_type == 'excel':
            return self.read_excel_data(file_path)
        elif file_type == 'word':
            return self.read_word_data(file_path)
        elif file_type == 'markdown':
            return self.read_markdown_data(file_path)
        elif file_type == 'text':
            return self.read_text_data(file_path)
        else:
            return {'error': f'不支持的文件类型: {file_path}'}
    
    def read_excel_template(self, file_path: str) -> Dict[str, Any]:
        """读取Excel模板"""
        result = {
            'path': file_path,
            'type': 'excel',
            'columns': [],
            'data': None
        }
        
        try:
            df = pd.read_excel(file_path)
            result['columns'] = list(df.columns)
            result['data'] = df
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_word_template(self, file_path: str) -> Dict[str, Any]:
        """读取Word模板"""
        result = {
            'path': file_path,
            'type': 'word',
            'columns': [],
            'tables': []
        }
        
        try:
            doc = Document(file_path)
            for table in doc.tables:
                table_info = {
                    'headers': [cell.text.strip() for cell in table.rows[0].cells],
                    'rows': len(table.rows),
                    'cols': len(table.columns)
                }
                result['tables'].append(table_info)
            
            if result['tables']:
                result['columns'] = result['tables'][0]['headers']
                
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def read_template(self, file_path: str) -> Dict[str, Any]:
        """读取模板 - 自动识别文件类型"""
        file_type = self.get_file_type(file_path)
        
        print(f"读取模板: {os.path.basename(file_path)}")
        print(f"  模板类型: {file_type}")
        
        if file_type == 'excel':
            return self.read_excel_template(file_path)
        elif file_type == 'word':
            return self.read_word_template(file_path)
        else:
            return {'error': f'不支持的模板类型: {file_path}'}
    
    def match_columns(self, template_cols: List[str], data_cols: List[str]) -> Dict[str, str]:
        """匹配模板列和数据列"""
        mapping = {}
        used_data_cols = set()
        
        sorted_template_cols = sorted(template_cols, key=len, reverse=True)
        
        for t_col in sorted_template_cols:
            best_match = None
            best_score = 0
            
            for d_col in data_cols:
                if d_col in used_data_cols:
                    continue
                
                score = self.similarity(t_col, d_col)
                
                if score < 0.6:
                    t_keywords = set(re.findall(r'[\u4e00-\u9fa5a-zA-Z]+', t_col))
                    d_keywords = set(re.findall(r'[\u4e00-\u9fa5a-zA-Z]+', d_col))
                    if t_keywords & d_keywords:
                        score = max(score, 0.7)
                
                if score > best_score:
                    best_score = score
                    best_match = d_col
            
            if best_match and best_score >= 0.5:
                mapping[t_col] = best_match
                used_data_cols.add(best_match)
        
        return mapping
    
    def apply_filters(self, df: pd.DataFrame, params: Dict[str, Any]) -> pd.DataFrame:
        """应用过滤条件"""
        if not params:
            return df
        
        result_df = df.copy()
        
        if params.get('date_range'):
            date_cols = [col for col in df.columns if '日期' in col or '时间' in col or 'date' in col.lower()]
            if date_cols:
                date_col = date_cols[0]
                result_df[date_col] = pd.to_datetime(result_df[date_col], errors='coerce')
                start, end = params['date_range']
                result_df = result_df[(result_df[date_col] >= start) & (result_df[date_col] <= end)]
        
        for filter_key, filter_values in params.get('filters', {}).items():
            matching_cols = [col for col in df.columns if filter_key in col]
            if matching_cols:
                col = matching_cols[0]
                result_df = result_df[result_df[col].isin(filter_values)]
        
        return result_df
    
    def extract_data_from_text_intelligent(self, text: str, template_cols: List[str]) -> List[Dict]:
        """智能文本数据提取"""
        data_rows = []
        
        has_city_col = any('城市' in col or '城市名' in col for col in template_cols)
        
        if has_city_col:
            data_rows = self._extract_city_economic_data(text, template_cols)
            if data_rows:
                return data_rows
        
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        current_record = {}
        
        for para in paragraphs:
            numbers = self.extract_numbers_from_text(para)
            
            for col in template_cols:
                col_keywords = re.findall(r'[\u4e00-\u9fa5a-zA-Z]+', col)
                
                for keyword in col_keywords:
                    if keyword in para and numbers:
                        if col not in current_record:
                            idx = para.find(keyword)
                            if idx != -1:
                                nearby_text = para[max(0, idx-10):min(len(para), idx+50)]
                                nearby_numbers = self.extract_numbers_from_text(nearby_text)
                                if nearby_numbers:
                                    current_record[col] = nearby_numbers[0]
                        break
            
            if len(current_record) >= max(1, len(template_cols) // 2):
                data_rows.append(current_record.copy())
                current_record = {}
        
        if current_record:
            data_rows.append(current_record)
        
        for row in data_rows:
            for col in template_cols:
                if col not in row:
                    row[col] = None
        
        return data_rows
    
    def _extract_city_economic_data(self, text: str, template_cols: List[str]) -> List[Dict]:
        """从城市经济报告文本中提取数据"""
        data_rows = []
        
        city_keywords = ['上海', '北京', '深圳', '重庆', '广州', '苏州', '杭州', 
                        '南京', '武汉', '成都', '西安', '郑州', '长沙', '青岛',
                        '东莞', '佛山', '宁波', '无锡', '合肥', '福州', '厦门',
                        '济南', '大连', '沈阳', '哈尔滨', '长春', '石家庄', '太原',
                        '南昌', '南宁', '昆明', '贵阳', '兰州', '海口', '呼和浩特',
                        '银川', '西宁', '拉萨', '乌鲁木齐', '南通', '泉州', '常州',
                        '烟台', '绍兴', '台州', '嘉兴', '潍坊', '扬州', '镇江',
                        '泰州', '济宁', '徐州', '惠州', '珠海', '中山', '盐城',
                        '临沂', '淄博', '江门', '威海', '菏泽', '湖州', '包头',
                        '洛阳', '沧州', '榆林', '岳阳', '株洲', '衡阳', '漳州',
                        '芜湖', '廊坊', '赣州']
        
        paragraphs = text.split('\n')
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            record = {}
            
            city_name = None
            for city in city_keywords:
                if city in para:
                    city_name = city
                    break
            
            if not city_name:
                continue
            
            for col in template_cols:
                if '城市' in col:
                    record[col] = city_name
                elif 'GDP' in col and '人均' not in col:
                    match = re.search(r'([0-9,\.]+)\s*亿[^0-9]*GDP|GDP[^0-9]*([0-9,\.]+)\s*亿', para)
                    if match:
                        record[col] = (match.group(1) or match.group(2)).replace(',', '')
                elif '人均GDP' in col or '人均' in col:
                    match = re.search(r'人均\s*GDP[^0-9]*([0-9,\.]+)', para)
                    if match:
                        record[col] = match.group(1).replace(',', '')
                elif '人口' in col:
                    match = re.search(r'([0-9,\.]+)\s*万[^0-9]*(?:常住)?人口|常住人口[^0-9]*([0-9,\.]+)\s*万', para)
                    if match:
                        record[col] = (match.group(1) or match.group(2)).replace(',', '')
                elif '预算' in col or '财政' in col:
                    match = re.search(r'预算收入[^0-9]*([0-9,\.]+)\s*亿', para)
                    if match:
                        record[col] = match.group(1).replace(',', '')
            
            if len(record) > 1:
                data_rows.append(record)
        
        for row in data_rows:
            for col in template_cols:
                if col not in row:
                    row[col] = None
        
        return data_rows
    
    def fill_excel_template(self, data_source: Dict, template: Dict, output_path: str, params: Dict = None) -> Dict:
        """填充Excel模板 - 支持所有数据源类型"""
        print(f"\n{'='*70}")
        print("填充Excel模板")
        print(f"数据源类型: {data_source['type']}")
        print('='*70)
        
        template_cols = template['columns']
        print(f"模板列: {template_cols}")
        
        if data_source['type'] == 'excel' and data_source.get('data') is not None:
            df = data_source['data'].copy()
            print(f"数据源行数: {len(df)}")
            
            if params:
                df = self.apply_filters(df, params)
                print(f"过滤后行数: {len(df)}")
            
            col_mapping = self.match_columns(template_cols, list(df.columns))
            print(f"\n列映射:")
            for t_col, d_col in col_mapping.items():
                score = self.similarity(t_col, d_col)
                print(f"  {t_col} <- {d_col} (相似度: {score:.2f})")
            
            result_data = []
            for _, row in df.iterrows():
                result_row = {}
                for t_col in template_cols:
                    if t_col in col_mapping:
                        source_col = col_mapping[t_col]
                        value = row.get(source_col)
                        result_row[t_col] = value if pd.notna(value) else None
                    else:
                        result_row[t_col] = None
                result_data.append(result_row)
            
            result_df = pd.DataFrame(result_data)
            
        elif data_source['type'] == 'word':
            if data_source.get('data') is not None:
                df = data_source['data'].copy()
                if params:
                    df = self.apply_filters(df, params)
                
                col_mapping = self.match_columns(template_cols, list(df.columns))
                
                result_data = []
                for _, row in df.iterrows():
                    result_row = {}
                    for t_col in template_cols:
                        if t_col in col_mapping:
                            result_row[t_col] = row.get(col_mapping[t_col])
                        else:
                            result_row[t_col] = None
                    result_data.append(result_row)
                
                result_df = pd.DataFrame(result_data)
                print(f"从Word表格提取: {len(result_df)}行")
            else:
                extracted_data = self.extract_data_from_text_intelligent(data_source.get('text', ''), template_cols)
                result_df = pd.DataFrame(extracted_data)
                print(f"从Word文本智能提取: {len(result_df)}行")
                
        elif data_source['type'] == 'markdown':
            if data_source.get('data') is not None:
                df = data_source['data'].copy()
                if params:
                    df = self.apply_filters(df, params)
                
                col_mapping = self.match_columns(template_cols, list(df.columns))
                
                result_data = []
                for _, row in df.iterrows():
                    result_row = {}
                    for t_col in template_cols:
                        if t_col in col_mapping:
                            result_row[t_col] = row.get(col_mapping[t_col])
                        else:
                            result_row[t_col] = None
                    result_data.append(result_row)
                
                result_df = pd.DataFrame(result_data)
                print(f"从Markdown表格提取: {len(result_df)}行")
            else:
                extracted_data = self.extract_data_from_text_intelligent(data_source.get('text', ''), template_cols)
                result_df = pd.DataFrame(extracted_data)
                print(f"从Markdown文本智能提取: {len(result_df)}行")
                
        elif data_source['type'] == 'text':
            if data_source.get('data') is not None:
                df = data_source['data'].copy()
                if params:
                    df = self.apply_filters(df, params)
                
                col_mapping = self.match_columns(template_cols, list(df.columns))
                
                result_data = []
                for _, row in df.iterrows():
                    result_row = {}
                    for t_col in template_cols:
                        if t_col in col_mapping:
                            result_row[t_col] = row.get(col_mapping[t_col])
                        else:
                            result_row[t_col] = None
                    result_data.append(result_row)
                
                result_df = pd.DataFrame(result_data)
                print(f"从文本表格提取: {len(result_df)}行")
            else:
                extracted_data = self.extract_data_from_text_intelligent(data_source.get('text', ''), template_cols)
                result_df = pd.DataFrame(extracted_data)
                print(f"从文本智能提取: {len(result_df)}行")
        
        else:
            result_df = pd.DataFrame(columns=template_cols)
        
        for col in result_df.columns:
            result_df[col] = result_df[col].fillna('N/A')
        
        result_df.to_excel(output_path, index=False)
        
        total_cells = len(result_df) * len(result_df.columns)
        filled_cells = total_cells - (result_df == 'N/A').sum().sum()
        fill_rate = (filled_cells / total_cells * 100) if total_cells > 0 else 0
        
        print(f"\n填充统计:")
        print(f"  总行数: {len(result_df)}")
        print(f"  填充率: {fill_rate:.1f}%")
        print(f"  输出: {output_path}")
        
        return {
            'output_path': output_path,
            'rows': len(result_df),
            'fill_rate': fill_rate
        }
    
    def fill_word_template(self, data_source: Dict, template: Dict, output_path: str, params: Dict = None) -> Dict:
        """填充Word模板 - 支持所有数据源类型"""
        print(f"\n{'='*70}")
        print("填充Word模板")
        print(f"数据源类型: {data_source['type']}")
        print('='*70)
        
        doc = Document(template['path'])
        
        df = None
        if data_source['type'] == 'excel' and data_source.get('data') is not None:
            df = data_source['data'].copy()
        elif data_source['type'] == 'word' and data_source.get('data') is not None:
            df = data_source['data'].copy()
        elif data_source['type'] == 'markdown' and data_source.get('data') is not None:
            df = data_source['data'].copy()
        elif data_source['type'] == 'text' and data_source.get('data') is not None:
            df = data_source['data'].copy()
        
        if df is None:
            print("数据源不包含表格数据，尝试从文本智能提取...")
            df = self._extract_dataframe_from_text(data_source, doc)
            if df is None or len(df) == 0:
                print("警告: 无法从文本提取数据")
                return {'output_path': output_path, 'fill_rate': 0}
        
        print(f"数据源行数: {len(df)}")
        
        if params:
            df = self.apply_filters(df, params)
            print(f"过滤后行数: {len(df)}")
        
        filter_values = []
        if params and params.get('filters'):
            for values in params['filters'].values():
                filter_values.extend(values)
        
        total_cells = 0
        filled_cells = 0
        
        for table_idx, table in enumerate(doc.tables):
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            target_value = filter_values[table_idx] if table_idx < len(filter_values) else None
            
            if target_value:
                matching_cols = [col for col in df.columns if any(v in col for v in ['城市', '地区', '名称'])]
                if matching_cols:
                    table_data = df[df[matching_cols[0]] == target_value]
                else:
                    table_data = df
            else:
                table_data = df
            
            print(f"表格{table_idx + 1}: {len(table_data)}行数据")
            
            col_mapping = self.match_columns(headers, list(df.columns))
            
            for row_idx in range(1, len(table.rows)):
                row = table.rows[row_idx]
                data_row_idx = row_idx - 1
                
                for col_idx, header in enumerate(headers):
                    if col_idx >= len(row.cells):
                        break
                    
                    total_cells += 1
                    
                    if data_row_idx < len(table_data):
                        data_row = table_data.iloc[data_row_idx]
                        
                        if header in col_mapping:
                            source_col = col_mapping[header]
                            value = data_row.get(source_col)
                            if pd.notna(value):
                                row.cells[col_idx].text = str(value)
                                filled_cells += 1
                            else:
                                row.cells[col_idx].text = "N/A"
                        else:
                            row.cells[col_idx].text = "N/A"
                    else:
                        row.cells[col_idx].text = "N/A"
        
        doc.save(output_path)
        
        fill_rate = (filled_cells / total_cells * 100) if total_cells > 0 else 0
        print(f"\n填充统计:")
        print(f"  总单元格: {total_cells}")
        print(f"  已填充: {filled_cells}")
        print(f"  填充率: {fill_rate:.1f}%")
        print(f"  输出: {output_path}")
        
        return {
            'output_path': output_path,
            'fill_rate': fill_rate
        }
    
    def _extract_dataframe_from_text(self, data_source: Dict, doc: Document) -> Optional[pd.DataFrame]:
        """从文本数据源提取DataFrame用于填充Word模板"""
        text = data_source.get('text', '')
        if not text:
            return None
        
        all_headers = set()
        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            all_headers.update(headers)
        
        if not all_headers:
            return None
        
        print(f"模板需要的列: {all_headers}")
        
        extracted_data = self.extract_data_from_text_intelligent(text, list(all_headers))
        
        if extracted_data:
            df = pd.DataFrame(extracted_data)
            print(f"从文本智能提取: {len(df)}行数据")
            return df
        
        return None
    
    def process(self, data_path: str, template_path: str, requirement: str = None, output_path: str = None) -> Dict:
        """通用处理入口 - 支持所有文件类型组合"""
        print(f"\n{'='*70}")
        print(f"处理任务")
        print(f"数据源: {os.path.basename(data_path)}")
        print(f"模板: {os.path.basename(template_path)}")
        print('='*70)
        
        params = self.parse_requirement(requirement)
        print(f"\n解析参数:")
        if params['date_range']:
            print(f"  日期范围: {params['date_range'][0]} - {params['date_range'][1]}")
        if params['filters']:
            print(f"  过滤条件: {params['filters']}")
        if params['keywords']:
            print(f"  关键词: {params['keywords']}")
        
        data_source = self.read_data_source(data_path)
        if 'error' in data_source:
            raise ValueError(f"数据源读取失败: {data_source['error']}")
        
        template = self.read_template(template_path)
        if 'error' in template:
            raise ValueError(f"模板读取失败: {template['error']}")
        
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            data_name = Path(data_path).stem
            template_ext = Path(template_path).suffix
            output_path = f"results/{data_name}_output_{timestamp}{template_ext}"
        
        os.makedirs(os.path.dirname(output_path) or 'results', exist_ok=True)
        
        if template['type'] == 'excel':
            return self.fill_excel_template(data_source, template, output_path, params)
        elif template['type'] == 'word':
            return self.fill_word_template(data_source, template, output_path, params)
        else:
            raise ValueError(f"不支持的模板类型: {template['type']}")


def main():
    parser = argparse.ArgumentParser(description='通用数据填充工具 - 支持Excel/Word/Markdown/Text全组合')
    parser.add_argument('--data', '-d', required=True, help='数据源文件路径 (支持.xlsx/.xls/.docx/.md/.txt)')
    parser.add_argument('--template', '-t', required=True, help='模板文件路径 (支持.xlsx/.xls/.docx)')
    parser.add_argument('--requirement', '-r', default='', help='用户要求')
    parser.add_argument('--output', '-o', default='', help='输出文件路径')
    parser.add_argument('--config', '-c', default='', help='配置文件路径(JSON)')
    
    args = parser.parse_args()
    
    config = None
    if args.config and os.path.exists(args.config):
        with open(args.config, 'r', encoding='utf-8') as f:
            config = json.load(f)
    
    filler = UniversalDataFiller(config)
    result = filler.process(
        data_path=args.data,
        template_path=args.template,
        requirement=args.requirement,
        output_path=args.output if args.output else None
    )
    
    print(f"\n✓ 处理完成!")
    print(f"  输出文件: {result['output_path']}")
    print(f"  填充率: {result['fill_rate']:.1f}%")


if __name__ == "__main__":
    main()
